"""
Generate competitive analysis Word document for GeoAsset Intelligence startup.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────────────────────────────────────
# COLOUR PALETTE
# ─────────────────────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x0D, 0x2B, 0x4E)   # #0D2B4E
MID_BLUE    = RGBColor(0x1A, 0x5C, 0x8E)   # #1A5C8E
ACCENT_BLUE = RGBColor(0x2E, 0x86, 0xC1)   # #2E86C1
LIGHT_GREY  = RGBColor(0xF4, 0xF6, 0xF7)   # #F4F6F7
MID_GREY    = RGBColor(0xAB, 0xB2, 0xB9)   # #ABB2B9
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x17, 0x20, 0x2A)   # near-black
GREEN       = RGBColor(0x1E, 0x8B, 0x4C)   # dark green
ORANGE      = RGBColor(0xE6, 0x7E, 0x22)   # orange
RED_SOFT    = RGBColor(0xC0, 0x39, 0x2B)   # soft red


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def rgb_hex(rgb: RGBColor) -> str:
    """Convert RGBColor to hex string (no #)."""
    return '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])


def set_cell_bg(cell, rgb: RGBColor):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(rgb))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Set border on a cell. kwargs: top, bottom, left, right each = dict(sz, color, val)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        if side in kwargs:
            border = OxmlElement(f'w:{side}')
            opts = kwargs[side]
            border.set(qn('w:val'), opts.get('val', 'single'))
            border.set(qn('w:sz'), str(opts.get('sz', 4)))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), opts.get('color', 'auto'))
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, size=10,
            color=BLACK, font_name='Calibri'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font_name
    return run


def heading(doc, text, level=1):
    """Add a styled heading."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(4)
    if level == 1:
        run = para.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = DARK_BLUE
        run.font.name = 'Calibri'
        # Bottom border
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), rgb_hex(ACCENT_BLUE))
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = MID_BLUE
        run.font.name = 'Calibri'
    elif level == 3:
        run = para.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = ACCENT_BLUE
        run.font.name = 'Calibri'
    return para


def body(doc, text, indent=0):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Inches(indent * 0.25)
    add_run(para, text, size=10, color=BLACK)
    return para


def bullet(doc, text, level=0, bold_prefix=None):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    if bold_prefix:
        add_run(para, bold_prefix + ': ', bold=True, size=10, color=DARK_BLUE)
        add_run(para, text, size=10, color=BLACK)
    else:
        add_run(para, text, size=10, color=BLACK)
    return para


def spacer(doc, height_pt=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    run.font.size = Pt(height_pt)
    return para


def section_intro(doc, text):
    """Lightly styled intro paragraph in italic grey."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(8)
    add_run(para, text, italic=True, size=10, color=MID_GREY)
    return para


def page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type())
    return para


def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    return br  # we'll handle this differently


def add_page_break(doc):
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# COMPANY CARD TABLE
# ─────────────────────────────────────────────────────────────────────────────
def company_card(doc, company):
    """
    company = {
        name, founded, hq, stage,
        core_product, target_customers,
        funding_revenue, differentiators (list), limitations (list),
        relevance  # str: Low / Medium / High
    }
    """
    heading(doc, company['name'], level=2)

    # meta line
    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(1)
    meta.paragraph_format.space_after = Pt(6)
    add_run(meta, f"Founded: {company.get('founded','N/A')}  |  ", bold=True, size=9.5, color=MID_GREY)
    add_run(meta, f"HQ: {company.get('hq','N/A')}  |  ", size=9.5, color=MID_GREY)
    add_run(meta, f"Stage/Type: {company.get('stage','N/A')}  |  ", size=9.5, color=MID_GREY)
    rel = company.get('relevance', 'Medium')
    rel_color = GREEN if rel == 'High' else (ORANGE if rel == 'Medium' else MID_GREY)
    add_run(meta, f"Competitive Relevance: ", bold=True, size=9.5, color=MID_GREY)
    add_run(meta, rel, bold=True, size=9.5, color=rel_color)

    # 2-column table for details
    tbl = doc.add_table(rows=5, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'

    col_labels = [
        'Core Product / Service',
        'Target Customers',
        'Funding / Revenue',
        'Key Differentiators',
        'Limitations / Gaps',
    ]
    col_values = [
        company.get('core_product', ''),
        company.get('target_customers', ''),
        company.get('funding_revenue', ''),
        '\n'.join(f'• {d}' for d in company.get('differentiators', [])),
        '\n'.join(f'• {l}' for l in company.get('limitations', [])),
    ]

    for i, (label, value) in enumerate(zip(col_labels, col_values)):
        row = tbl.rows[i]
        # label cell
        lc = row.cells[0]
        set_cell_bg(lc, DARK_BLUE)
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(lp, label, bold=True, size=9.5, color=WHITE)
        lc.width = Inches(1.8)

        # value cell
        vc = row.cells[1]
        set_cell_bg(vc, LIGHT_GREY)
        vp = vc.paragraphs[0]
        add_run(vp, value, size=9.5, color=BLACK)
        vc.width = Inches(4.5)

    spacer(doc, 10)


# ─────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
def add_cover(doc):
    # Title block
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(60)
    para.paragraph_format.space_after = Pt(6)
    add_run(para, 'COMPETITIVE ANALYSIS', bold=True, size=24,
            color=DARK_BLUE, font_name='Calibri')

    para2 = doc.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2.paragraph_format.space_before = Pt(4)
    para2.paragraph_format.space_after = Pt(4)
    add_run(para2,
            'Geospatial AI & Physical Asset Intelligence',
            bold=True, size=17, color=ACCENT_BLUE, font_name='Calibri')

    para3 = doc.add_paragraph()
    para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para3.paragraph_format.space_before = Pt(2)
    para3.paragraph_format.space_after = Pt(40)
    add_run(para3,
            'Market Landscape for a GeoAsset Intelligence Startup',
            italic=True, size=12, color=MID_GREY, font_name='Calibri')

    # Divider
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div.paragraph_format.space_before = Pt(0)
    div.paragraph_format.space_after = Pt(20)
    add_run(div, '─' * 60, size=10, color=ACCENT_BLUE)

    # Meta
    for line in ['Prepared for: Internal Strategy', 'Date: May 2026', 'Confidential']:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.paragraph_format.space_before = Pt(2)
        add_run(mp, line, size=10, color=MID_GREY, font_name='Calibri')

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# EXECUTIVE SUMMARY
# ─────────────────────────────────────────────────────────────────────────────
def add_exec_summary(doc):
    heading(doc, 'Executive Summary', level=1)
    section_intro(doc,
        'This document provides a structured competitive analysis of the emerging market '
        'for geospatial AI and physical asset intelligence, with particular focus on the '
        'intersection of satellite/remote-sensing data, AI analytics, productive assets '
        '(industrial, energy, mining, real estate, agriculture), and ESG/financial use cases.')

    heading(doc, 'Market Context', level=2)
    body(doc,
        'The global geospatial AI market was valued at ~USD 38 billion in 2024 and is '
        'projected to reach USD 64.6 billion by 2030 (CAGR ~9%). The satellite-based '
        'Earth observation segment alone stood at USD 3.7 billion in 2024. Venture '
        'investment in AI broadly surpassed USD 200 billion in 2025, with physical-world '
        'and geospatial AI attracting a growing share.')

    heading(doc, 'Key Findings', level=2)
    findings = [
        ('Market fragmentation', 'No single player dominates the full stack from raw '
         'satellite data to ESG/financial decision-support for physical productive assets. '
         'Most incumbents own one layer (data OR analytics OR ESG ratings).'),
        ('Satellite imagery leaders are being acquired', 'Orbital Insight (acq. Privateer 2024), '
         'Descartes Labs (acq. EarthDaily 2024), Geosite (acq. Descartes Labs 2024), '
         'Rezatec (majority stake 2024) signal consolidation.'),
        ('ESG data is disconnected from physical location', 'MSCI, Sustainalytics, and S&P Trucost '
         'rate companies, not assets. Their physical risk datasets map company HQ or '
         'disclosed facility lists — they do not continuously monitor actual asset activity.'),
        ('AI-native startups are gaining ground in narrow verticals', 'AiDash (utilities/vegetation), '
         'LiveEO (infrastructure), Kayrros (energy emissions), Satelytics (pipeline/utility) '
         'all raised significant rounds 2022-2025 but serve specific sectors.'),
        ('The white space', 'A platform that (1) locates & identifies productive assets globally '
         'using satellite + public data, (2) continuously monitors them via AI, and (3) '
         'links findings to ESG scores, financial risk, and supply-chain due diligence — '
         'at scale across sectors — does not yet exist as a standalone product.'),
    ]
    for bold_label, text in findings:
        bullet(doc, text, bold_prefix=bold_label)

    spacer(doc, 8)
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# CATEGORY 1: SATELLITE + AI
# ─────────────────────────────────────────────────────────────────────────────
def add_cat1(doc):
    heading(doc, 'Category 1: Satellite Imagery + AI Asset Intelligence', level=1)
    section_intro(doc,
        'These companies use multi-source satellite data (optical, SAR, hyperspectral) '
        'combined with AI/ML to detect, locate, and monitor physical assets and human '
        'activity on the ground. They primarily serve government, defence, commodities '
        'traders, and increasingly insurance and ESG teams.')

    companies = [
        {
            'name': 'Orbital Insight',
            'founded': '2013',
            'hq': 'Palo Alto, CA (acquired by Privateer, May 2024)',
            'stage': 'Acquired (prev. Series D)',
            'core_product':
                'Geospatial big data platform using AI to analyse satellite, UAV and IoT data. '
                'Products: supply-chain visibility, commodity tracking, geopolitical intelligence.',
            'target_customers': 'Government agencies, hedge funds, commodities traders, NGOs',
            'funding_revenue': 'Total raised: ~$132M (Series D $50M, Sequoia-led, 2019). '
                               'Revenue ~$44M (June 2024 data). Acquired by Privateer for undisclosed sum, May 2024.',
            'differentiators': [
                'Pioneer in geospatial AI — deep IP in change-detection at scale',
                'Strong US government & IC relationships',
                'Multi-source fusion (optical + SAR + IoT)',
            ],
            'limitations': [
                'Brand/product continuity uncertain post-acquisition by Privateer (Wozniak space-sustainability focus)',
                'Historically weak on ESG and financial analytics layer',
                'Largely US-centric customer base',
            ],
            'relevance': 'High',
        },
        {
            'name': 'Kayrros',
            'founded': '2016',
            'hq': 'Paris, France',
            'stage': 'Series C (acquired by Energy Aspects, March 2026)',
            'core_product':
                'Environmental intelligence platform tracking energy assets, methane/GHG emissions, '
                'and climate indicators via satellite + AI. Products include methane monitoring, '
                'oil/gas storage tracking, wildfire risk, and renewable energy asset intelligence.',
            'target_customers': 'Energy companies, commodity traders, regulators, ESG investors, banks',
            'funding_revenue': 'Total raised: ~$78M (Series C $44M, 2022). Revenue ~$20.2M (Oct 2024, 80 customers). '
                               'Acquired by Energy Aspects (March 2026). Named Fortune "Change the World" 2024 #30.',
            'differentiators': [
                'Best-in-class methane/GHG emissions tracking from space — regulatory compliance use case',
                'Near-real-time energy asset monitoring (tanks, refineries, LNG terminals)',
                'Strong traction in financial/ESG community',
                'TIME Best Inventions 2025 (Wildfire Risk Monitor)',
            ],
            'limitations': [
                'Revenue still relatively small (~$20M); limited outside energy sector',
                'Post-acquisition direction may shift toward Energy Aspects clients only',
                'Does not cover mining, manufacturing, or agri assets',
            ],
            'relevance': 'High',
        },
        {
            'name': 'Planet Labs',
            'founded': '2010',
            'hq': 'San Francisco, CA (NASDAQ: PL)',
            'stage': 'Public (NYSE: PL)',
            'core_product':
                'World\'s largest commercial satellite constellation (200+ small sats). Provides '
                'daily optical imagery of entire Earth surface. Analytics APIs and tasking services.',
            'target_customers': 'Governments, agriculture, forestry, insurance, mining, infrastructure',
            'funding_revenue': 'FY2024 revenue: ~$244M; FY2026 guidance $281–289M (+14–18% YoY). '
                               'USD 280M contract with Germany (2025). Market cap ~$1B range.',
            'differentiators': [
                'Unmatched revisit frequency (daily global coverage)',
                'Archive depth — decades of imagery for change detection',
                'Azure/Microsoft integration for enterprise distribution',
                'Broad sector coverage',
            ],
            'limitations': [
                'Primarily a data/infrastructure provider — limited vertical analytics',
                'Not profitable; high capex for constellation maintenance',
                'Customers must build their own AI/analytics layer or use 3rd parties',
                'Optical only — limited in cloud cover, night-time, covert activity',
            ],
            'relevance': 'Medium',
        },
        {
            'name': 'ICEYE',
            'founded': '2014',
            'hq': 'Espoo, Finland',
            'stage': 'Late-stage private (Series E+)',
            'core_product':
                'World\'s largest SAR (synthetic aperture radar) satellite constellation (54+ sats). '
                'All-weather, day/night persistent monitoring. Key applications: insurance, '
                'disaster response, defence, utilities (grid monitoring), flood mapping.',
            'target_customers': 'Insurance/reinsurance, governments, defence, utilities, financial institutions',
            'funding_revenue': 'Total raised: ~$430M+. Dec 2025: EUR 150M round, valuation EUR 2.4B (~USD 2.8B). '
                               'Jan 2024: $65M additional; 2024 total raises ~$158M.',
            'differentiators': [
                'SAR works through clouds, at night — critical for persistent asset monitoring',
                'Three revenue streams: imagery, analytical services, constellation-as-a-service',
                'Growing insurance vertical (flood/hurricane) with quantified asset impact',
                'European-sovereign capability — strong non-US positioning',
            ],
            'limitations': [
                'Analytics still at early stage for most sectors outside insurance/defence',
                'SAR imagery is harder to interpret than optical — limits self-serve customers',
                'ESG integration not a primary product focus',
            ],
            'relevance': 'Medium',
        },
        {
            'name': 'Satellogic',
            'founded': '2010',
            'hq': 'Montevideo, Uruguay / Miami, FL (NASDAQ: SATL)',
            'stage': 'Public (NASDAQ: SATL)',
            'core_product':
                'AI-first satellite constellation providing sub-meter resolution optical imagery '
                'with on-orbit AI analytics. Target: near-daily tasking, edge AI processing. '
                '2024 revenue: $12.9M (+28% YoY).',
            'target_customers': 'Government, defence, intelligence, commercial analytics',
            'funding_revenue': '$12.9M revenue (2024, +28% YoY). Nasdaq-listed; $30M defence contract 2025. '
                               'Liquidity challenges noted by analysts.',
            'differentiators': [
                'On-orbit AI processing — reduces latency dramatically',
                'Very high resolution (30cm class) with frequent revisit',
                'NASA Commercial SmallSat Data Acquisition Program selection (2024)',
            ],
            'limitations': [
                'Small commercial revenue; heavily government/defence dependent',
                'Liquidity concerns cited by analysts in 2025',
                'No ESG or financial analytics products',
                'NextGen 30cm satellite not operational until ~2027',
            ],
            'relevance': 'Low',
        },
        {
            'name': 'Descartes Labs (now EarthDaily Analytics)',
            'founded': '2014',
            'hq': 'Santa Fe, NM (acquired by EarthDaily, Oct 2024)',
            'stage': 'Acquired',
            'core_product':
                'Geospatial intelligence platform with AI analytics for commodity intelligence '
                '(Marigold), insurance risk (Iris), and mineral exploration (Ascend). '
                'Acquired Geosite (Apr 2024) for insurance geospatial data.',
            'target_customers': 'Insurance, commodity trading, mining/minerals, agriculture',
            'funding_revenue': 'Total raised: ~$58.4M pre-acquisition. Acquired by EarthDaily (Oct 2024) '
                               'ahead of EarthDaily constellation launch in 2025.',
            'differentiators': [
                'Vertical AI products (not raw data) — Marigold, Iris, Ascend',
                'Strong commodity/insurance analytics heritage',
                'Now coupled with EarthDaily\'s new hyperspectral constellation',
            ],
            'limitations': [
                'Integration risk post-acquisition',
                'EarthDaily constellation deployment delayed',
                'No dedicated ESG or real-time industrial monitoring product',
            ],
            'relevance': 'Medium',
        },
        {
            'name': 'Ursa Space',
            'founded': '2015',
            'hq': 'Ithaca, NY',
            'stage': 'Growth (total raised ~$69.7M)',
            'core_product':
                'SAR-based satellite intelligence infrastructure. Provides SAR image catalog '
                'and analytics for change detection, maritime, energy assets, and defence. '
                'Key product: persistent monitoring of industrial/energy facilities.',
            'target_customers': 'Government, defence/intelligence, energy companies, financial analysts',
            'funding_revenue': 'Total raised: ~$69.7M. Aug 2025: Corporate Minority round. '
                               '$151B SHIELD IDIQ contract ceiling (MDA, Dec 2025).',
            'differentiators': [
                'SAR specialist — all-weather persistent monitoring of ground assets',
                'Energy facility tracking (storage, production assets)',
                'Strong government/IC contracts provide revenue floor',
            ],
            'limitations': [
                'Primarily government/defence revenue — limited commercial ESG/finance',
                'Small team; limited geographic sales coverage',
                'No integrated ESG or financial scoring layer',
            ],
            'relevance': 'Medium',
        },
        {
            'name': 'SpaceKnow',
            'founded': '2014',
            'hq': 'Prague, Czech Republic',
            'stage': 'Early-stage (Series A, $5.45M — last round 2017)',
            'core_product':
                'Satellite data analytics platform using AI for industrial monitoring, '
                'economic activity tracking, defence, and environmental monitoring. '
                'Known for China PMI / manufacturing activity indices.',
            'target_customers': 'Financial analysts, hedge funds, government, intelligence',
            'funding_revenue': 'Total raised: $5.45M (Series A, 2017). No major recent rounds.',
            'differentiators': [
                'Economic activity indices from satellite imagery — used by macro investors',
                'Early mover in satellite-based financial analytics',
                'Multi-satellite data fusion',
            ],
            'limitations': [
                'Severely underfunded vs peers',
                'No significant product updates visible since 2022',
                'Narrow use case (macro economic indicators)',
                'Risk of being outpaced by better-funded peers',
            ],
            'relevance': 'Low',
        },
        {
            'name': 'Satelytics',
            'founded': '2014',
            'hq': 'Bowling Green, OH',
            'stage': 'Early growth (total raised ~$5M)',
            'core_product':
                'Geospatial analytics SaaS for industrial sector. Uses AI + satellite/aerial '
                'imagery to detect environmental risks and operational anomalies in: oil & gas, '
                'pipelines, utilities, water, mining, forestry, specialty chemicals.',
            'target_customers': 'Oil & gas operators, pipeline companies, utilities, mining firms, municipalities',
            'funding_revenue': 'Total raised: ~$5M. Revenue: $6.1M (Oct 2024, +42% YoY from $4.3M). 35-person team.',
            'differentiators': [
                'Deep focus on environmental compliance and leak/spill detection',
                'Serves highly regulated industries with clear ROI (avoid fines/disasters)',
                'Broad multi-sector industrial coverage',
            ],
            'limitations': [
                'Very small — limited capacity to scale globally',
                'No ESG ratings or financial integration',
                'Reactive analytics (detect problems) vs proactive asset intelligence',
            ],
            'relevance': 'Medium',
        },
    ]

    for c in companies:
        company_card(doc, c)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# CATEGORY 2: PUBLIC / ALTERNATIVE DATA AGGREGATORS
# ─────────────────────────────────────────────────────────────────────────────
def add_cat2(doc):
    heading(doc, 'Category 2: Public & Alternative Data Aggregators for Asset Intelligence', level=1)
    section_intro(doc,
        'These are large, established data platforms that aggregate corporate, financial, '
        'regulatory, and in some cases physical-asset data. They serve financial institutions, '
        'risk managers, and corporate intelligence teams. Physical asset coverage is typically '
        'company-level (ownership records), not real-time asset-level monitoring.')

    companies = [
        {
            'name': 'S&P Global Market Intelligence',
            'founded': '1860 (S&P Global, Market Intelligence div. via SNL acquisition 2015)',
            'hq': 'New York, NY',
            'stage': 'Public (NYSE: SPGI)',
            'core_product':
                'Multi-asset data platform covering company financials, credit ratings, '
                'commodity data, and ESG. Includes Trucost (physical risk data for 2.9M+ '
                'assets mapped to 15,000+ listed companies across 7 climate hazards).',
            'target_customers': 'Banks, insurers, asset managers, corporates, governments',
            'funding_revenue': 'S&P Global FY2025 revenue: $15.3B (+8% YoY). '
                               'Market Intelligence segment: $4.92B (FY2025, +6% YoY). '
                               'Trucost: undisclosed segment, part of Sustainable1 division.',
            'differentiators': [
                'Largest physical climate risk dataset: 2.9M+ assets, 15K+ companies, 7 hazards',
                'Integrated with credit ratings — systemic financial risk linkage',
                'Trusted brand in financial institutions globally',
                'Distribution through Xpressfeed and direct API',
            ],
            'limitations': [
                'Physical risk is modelled/static — not real-time satellite monitoring',
                'Asset list depends on company self-disclosure + licence agreements',
                'Does not identify unknown/undisclosed assets',
                'ESG scores are lagging indicators, not operational intelligence',
                'Very expensive — pricing excludes startups and smaller institutions',
            ],
            'relevance': 'High',
        },
        {
            'name': 'MSCI ESG & GeoSpatial Asset Intelligence',
            'founded': '1969 (MSCI), GeoSpatial AI product launched ~2023-2024',
            'hq': 'New York, NY',
            'stage': 'Public (NYSE: MSCI)',
            'core_product':
                'ESG ratings for 8,500+ public companies, climate analytics, and (new) '
                'GeoSpatial Asset Intelligence: location-based data for ~70,000 public & '
                'private companies to assess physical hazard/nature risk at asset level.',
            'target_customers': 'Asset managers, banks, pension funds, corporates',
            'funding_revenue': 'MSCI FY2024 ESG analytics grew 40% YoY. '
                               'Total MSCI revenue ~$2.5B (2024). '
                               'Top 5 ESG vendors control ~65% of market.',
            'differentiators': [
                'GeoSpatial Asset Intelligence: 70K companies with location-linked physical risk — largest such dataset',
                'Integrated into existing MSCI portfolio analytics workflows',
                'Cross-use of Moody\'s private company/sovereign data (partnership)',
                'Multi-asset class: equities, fixed income, real assets',
            ],
            'limitations': [
                'Location data is largely company-reported/inferred — not satellite-verified',
                'No real-time operational monitoring of assets',
                'Designed for portfolio risk — not operational intelligence or due diligence',
                'Breadth at the expense of depth per asset',
            ],
            'relevance': 'High',
        },
        {
            'name': 'Morningstar Sustainalytics',
            'founded': '1992 (Sustainalytics, acquired by Morningstar 2020)',
            'hq': 'Amsterdam, Netherlands / Chicago, IL',
            'stage': 'Subsidiary of Morningstar (NASDAQ: MORN)',
            'core_product':
                'ESG Risk Ratings for 14,000+ companies. Corporate ESG research, '
                'second-party opinions (SPOs) on green bonds. Primarily company-level, '
                'not asset-level.',
            'target_customers': 'Asset managers, banks, pension funds, bond issuers',
            'funding_revenue': 'Morningstar FY2024 revenue: ~$2.0-2.2B. '
                               'Sustainalytics Q1 2025 revenue: $28.8M (-6.5% YoY) — declining. '
                               '22.7% share of SPO market (H1 2024). '
                               'Morningstar acquired DBRS ESG unit ($375M, Sep 2025).',
            'differentiators': [
                'Market-leading position in green/sustainable bond SPOs',
                'Broad institutional distribution via Morningstar platform',
                '14K+ company coverage',
            ],
            'limitations': [
                'Revenue declining — under pressure from AI-driven ESG data newcomers',
                'Exclusively company-level ratings, no asset-level intelligence',
                'Does not use satellite or geospatial data for monitoring',
                'Methodology opacity — facing regulatory scrutiny in EU',
            ],
            'relevance': 'Medium',
        },
        {
            'name': 'Dun & Bradstreet',
            'founded': '1841',
            'hq': 'Jacksonville, FL',
            'stage': 'Public (NYSE: DNB) — going private (Clearlake Capital, $7.7B deal, Mar 2025)',
            'core_product':
                'Business data platform: 500M+ business records, D-U-N-S numbers, '
                'firmographic intelligence, risk scores, supply chain transparency. '
                'No direct satellite monitoring.',
            'target_customers': 'Banks, corporates (procurement, risk, compliance), insurers',
            'funding_revenue': 'FY2024 revenue: $2.38B (+2.9% YoY). '
                               'Q1 2025 revenue: $579.8M (+2.7% YoY). '
                               'Being acquired by Clearlake for $7.7B (2025).',
            'differentiators': [
                '500M+ business records with D-U-N-S IDs — universal identifier layer',
                'Supply chain risk monitoring',
                'LSEG partnership for private markets data',
            ],
            'limitations': [
                'No geospatial or satellite-based asset monitoring',
                'Company-level data — does not track facility/asset-level activity',
                'Data quality concerns in emerging markets',
                'Growth stagnant (~3% YoY)',
            ],
            'relevance': 'Low',
        },
        {
            'name': 'Bureau van Dijk (Moody\'s)',
            'founded': '1991 (acquired by Moody\'s 2017)',
            'hq': 'Amsterdam, Netherlands',
            'stage': 'Subsidiary of Moody\'s (NYSE: MCO)',
            'core_product':
                'Orbis database: 400M+ company records globally, ownership structures, '
                'financial data, corporate hierarchies. Useful for identifying beneficial '
                'ownership of physical asset-holding companies.',
            'target_customers': 'Financial institutions, compliance teams, governments, consultancies',
            'funding_revenue': 'Part of Moody\'s Analytics; Moody\'s FY2024 revenue ~$7B. '
                               'BvD specific revenue not disclosed.',
            'differentiators': [
                'Deepest ownership/ultimate beneficial owner (UBO) data globally',
                'Useful for tracing who owns which industrial/physical assets via corporate chains',
                'Integration with Moody\'s credit analytics',
            ],
            'limitations': [
                'Purely company/ownership data — no asset-level monitoring',
                'Does not know what assets a company owns or where they are physically',
                'Static records — no real-time monitoring',
            ],
            'relevance': 'Low',
        },
    ]

    for c in companies:
        company_card(doc, c)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# CATEGORY 3: ESG & SUSTAINABILITY MONITORING
# ─────────────────────────────────────────────────────────────────────────────
def add_cat3(doc):
    heading(doc, 'Category 3: ESG & Sustainability Asset Monitoring Platforms', level=1)
    section_intro(doc,
        'These platforms focus on translating environmental, social, and governance data '
        'into actionable scores or insights for investors and regulators. Some use '
        'AI/alternative data but remain primarily focused on portfolio-level reporting '
        'rather than asset-level physical monitoring.')

    companies = [
        {
            'name': 'Clarity AI',
            'founded': '2017',
            'hq': 'Madrid, Spain / New York, NY',
            'stage': 'Late-stage private (Series D)',
            'core_product':
                'Sustainability tech platform using ML to provide ESG and climate insights. '
                'Covers 30,000+ organizations. Products: portfolio impact analysis, '
                'regulatory compliance (SFDR, EU Taxonomy), issuer research, '
                'sustainability AI research assistant (2025 launch).',
            'target_customers': 'Asset managers, banks, pension funds (clients manage >$60T AUM)',
            'funding_revenue': 'Total raised: ~$100-120M (sources vary; $12.9M round Sep 2025). '
                               'BlackRock minority investor. Named Forrester Wave leader Q3 2024. '
                               'Acquired ecolytiq (fintech climate engagement) Jul 2025.',
            'differentiators': [
                'AI-native — built with ML from ground up vs incumbent bolted-on analytics',
                'Regulatory compliance engine (SFDR, EU Taxonomy, CSRD)',
                'BlackRock distribution partnership',
                'Customisable — investors can weight their own materiality factors',
                'Named Forrester Wave ESG Data & Analytics leader (Q3 2024)',
            ],
            'limitations': [
                'No satellite or geospatial data for physical asset monitoring',
                'Company-level ESG scores, not asset-level monitoring',
                'Revenue scale not disclosed; likely still <$100M',
                'Dependent on company-reported data quality',
            ],
            'relevance': 'High',
        },
        {
            'name': 'Kayrros',
            'founded': '2016',
            'hq': 'Paris, France',
            'stage': 'Acquired by Energy Aspects (Mar 2026)',
            'core_product':
                '(See Category 1 for full profile.) Kayrros uniquely bridges Categories 1 and 3: '
                'satellite-derived GHG emissions and energy asset monitoring used directly for '
                'ESG reporting and regulatory compliance (EU Methane Regulation, SFDR).',
            'target_customers': 'Energy companies, ESG investors, commodity traders, regulators',
            'funding_revenue': 'See Category 1. Revenue ~$20.2M, raised ~$78M.',
            'differentiators': [
                'Satellite-derived emissions data — independent of company self-reporting',
                'Used for regulatory compliance — EUDR, EU Methane Reg.',
                'Fortune "Change the World" recognition',
            ],
            'limitations': [
                'Narrow energy/emissions focus',
                'Post-acquisition may narrow to energy clients only',
            ],
            'relevance': 'High',
        },
        {
            'name': 'MSCI ESG (Physical Risk / Trucost comparison)',
            'founded': '(See Category 2)',
            'hq': 'New York, NY',
            'stage': 'Public (NYSE: MSCI)',
            'core_product':
                'Physical climate risk scoring overlaid on asset locations. '
                'GeoSpatial Asset Intelligence product maps ~70K company locations to '
                'physical hazards (flood, fire, drought, etc). ESG ratings 8,500+ companies.',
            'target_customers': 'Asset managers, banks, pension funds',
            'funding_revenue': 'ESG segment +40% YoY 2024. Total MSCI revenue ~$2.5B.',
            'differentiators': [
                'Direct integration with portfolio management tools',
                'Broadest company coverage for physical climate risk',
            ],
            'limitations': [
                'Not real-time — climate risk modelled from static scenarios',
                'No satellite verification of actual asset conditions',
            ],
            'relevance': 'High',
        },
        {
            'name': 'S&P Global Trucost',
            'founded': '2000 (Trucost), acquired by S&P Global 2016',
            'hq': 'London, UK',
            'stage': 'Subsidiary of S&P Global',
            'core_product':
                'Environmental data for 15,000+ companies (99% global market cap). '
                'Physical Risk: 2.9M+ assets mapped to 7 climate hazards under multiple '
                'IPCC scenarios. Covers GHG, water use, biodiversity, fossil fuels.',
            'target_customers': 'Institutional investors, banks, insurers, corporates',
            'funding_revenue': 'Part of S&P Global ($15.3B FY2025 revenue). '
                               'Trucost distributed via Xpressfeed and Sustainable1.',
            'differentiators': [
                'Largest asset-level physical risk map (2.9M assets)',
                'Cross-referenced with financial exposure data',
                'State Street partnership — embedded in $40T AUM platform',
            ],
            'limitations': [
                'Physical risk is modelled/scenario-based, not satellite-observed',
                'Coverage relies on disclosed or inferred facility lists',
                'No operational monitoring — identifies risk exposure, not actual impact',
            ],
            'relevance': 'High',
        },
        {
            'name': 'AiDash',
            'founded': '2019',
            'hq': 'Santa Clara, CA',
            'stage': 'Series C',
            'core_product':
                'AI + satellite platform for utilities and infrastructure. Key product: '
                'Intelligent Vegetation Management System (IVMS) — uses satellite data '
                'and AI to manage vegetation encroachment on power lines and pipelines. '
                'Also: asset inspection, wildfire risk, predictive maintenance.',
            'target_customers': 'Electric utilities, gas utilities, telecom infrastructure operators',
            'funding_revenue': 'Total raised: $91.5M (Series C $58.5M, 2024 — oversubscribed). '
                               'Investors: Lightrock, National Grid Partners, Shell Ventures, G2VP.',
            'differentiators': [
                'Very strong product-market fit in utilities — clear cost/risk ROI',
                'Satellite + AI for continuous vegetation/asset monitoring',
                'Industrial investors (National Grid, Shell) validate use case',
                'Growing into European market with dedicated HQ',
            ],
            'limitations': [
                'Narrow vertical: primarily electric/gas utilities',
                'Does not produce ESG scores or financial analytics',
                'No coverage of manufacturing, mining, or agri assets',
            ],
            'relevance': 'High',
        },
    ]

    for c in companies:
        company_card(doc, c)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# CATEGORY 4: REAL ESTATE / INDUSTRIAL ASSET LOCATION
# ─────────────────────────────────────────────────────────────────────────────
def add_cat4(doc):
    heading(doc, 'Category 4: Real Estate & Industrial Asset Location Intelligence', level=1)
    section_intro(doc,
        'These platforms focus on locating, valuing, and analysing physical real estate '
        'or infrastructure assets. They use a mix of public records, proprietary surveys, '
        'and increasingly AI to provide decision intelligence for real estate, insurance, '
        'corporate facilities, and site selection.')

    companies = [
        {
            'name': 'CoStar Group',
            'founded': '1987',
            'hq': 'Washington, DC (NASDAQ: CSGP)',
            'stage': 'Public (NASDAQ: CSGP)',
            'core_product':
                'Commercial real estate data and analytics: property listings, valuations, '
                'lease comps, market analytics, LoopNet (commercial search), Apartments.com, '
                'Homes.com. 2024 acquisition of Matterport (3D property capture).',
            'target_customers': 'CRE brokers, property investors, lenders, occupiers, appraisers',
            'funding_revenue': 'FY2025 revenue: $3.2B (+19% YoY). FY2024 revenue: $2.74B (+11% YoY). '
                               'Market cap ~$30B+. Acquired Matterport and Domain (Australia).',
            'differentiators': [
                'Dominant in US commercial real estate — near-monopoly on CRE data',
                'Vertical integration from listing to financing analytics',
                'Matterport acquisition adds 3D digital twin capability',
            ],
            'limitations': [
                'Focused on real estate — not industrial/energy/mining assets',
                'No satellite monitoring or ESG analytics for physical assets',
                'Does not track operational activity of assets (just physical characteristics)',
                'Very US-centric',
            ],
            'relevance': 'Low',
        },
        {
            'name': 'Gordian (formerly RSMeans)',
            'founded': '1942 (RSMeans), Gordian formed 2017',
            'hq': 'Greenville, SC',
            'stage': 'Private (owned by PE)',
            'core_product':
                'Construction cost intelligence and facilities management. Products: '
                'RSMeans cost data, job order contracting (JOC) platform, facility '
                'condition assessments. Tracks building/infrastructure asset condition and costs.',
            'target_customers': 'Government agencies, universities, healthcare, facility managers',
            'funding_revenue': 'Revenue not publicly disclosed. PE-backed; primarily US public sector.',
            'differentiators': [
                'Dominant in US construction cost data (RSMeans)',
                'Asset condition lifecycle management',
            ],
            'limitations': [
                'Very niche — construction/facility management, not productive industrial assets',
                'No satellite, geospatial, or ESG capabilities',
                'Primarily US public sector',
            ],
            'relevance': 'Low',
        },
        {
            'name': 'Kroll',
            'founded': '1932',
            'hq': 'New York, NY',
            'stage': 'Private (PE-backed, Duff & Phelps merged 2018)',
            'core_product':
                'Valuation, due diligence, and risk advisory. Relevant products: '
                'fixed-asset appraisals, plant/machinery valuation, real estate appraisals, '
                'business intelligence investigations, environmental risk assessments.',
            'target_customers': 'Financial institutions, PE funds, corporates (M&A), insurers, litigation',
            'funding_revenue': 'Revenue not publicly disclosed. PE-backed (Jordan Company + others).',
            'differentiators': [
                'Trusted brand for asset valuation in M&A and litigation',
                'Combines financial + physical asset assessment in due diligence',
                'Global reach for cross-border asset intelligence',
            ],
            'limitations': [
                'Services-based, not a data/SaaS product',
                'No real-time or satellite-based monitoring',
                'High cost — advisory model, not scalable analytics',
            ],
            'relevance': 'Low',
        },
        {
            'name': 'Site Selection Group',
            'founded': '1997',
            'hq': 'Dallas, TX',
            'stage': 'Private (boutique)',
            'core_product':
                'Advisory firm for corporate site selection and real estate strategy. '
                'Uses GIS, demographics, incentive analysis, and labour market data to '
                'help corporations locate facilities (manufacturing, distribution, offices).',
            'target_customers': 'Corporates planning new facilities, economic development agencies',
            'funding_revenue': 'Small boutique; revenue not disclosed.',
            'differentiators': [
                'Specialised multi-criteria site selection methodology',
                'Strong incentive negotiation expertise',
            ],
            'limitations': [
                'Advisory only — no technology platform or SaaS product',
                'No continuous monitoring after site selection',
                'Not relevant for financial/ESG asset intelligence',
            ],
            'relevance': 'Low',
        },
        {
            'name': 'LiveEO',
            'founded': '2018',
            'hq': 'Berlin, Germany',
            'stage': 'Series B+',
            'core_product':
                'AI-powered satellite monitoring for civil infrastructure assets: power lines, '
                'railways, pipelines. Products: vegetation management, third-party interference '
                'detection, corridor mapping, climate risk. Launching Twinspector constellation '
                '(35cm stereo, 2028).',
            'target_customers': 'Electric utilities, railway operators, pipeline companies (EU + global)',
            'funding_revenue': 'Total raised: >€72M. Jun 2024: €25M Series B. 2026: >€28M first close new round. '
                               'Revenue 3x growth since 2021 ESA Demo Project.',
            'differentiators': [
                'Deep infrastructure asset monitoring — proven EU utility customers',
                'Developing own constellation (Twinspector) for independence',
                'Expanding into defence/security vertical',
                'European-built sovereign capability',
            ],
            'limitations': [
                'Infrastructure-focused — not manufacturing, mining, or agri assets',
                'No ESG scores or financial analytics layer',
                'Small compared to Maxar/Planet for global coverage',
            ],
            'relevance': 'High',
        },
        {
            'name': 'Rezatec',
            'founded': '2013',
            'hq': 'Oxford, UK',
            'stage': 'Acquired (majority stake, Dec 2024)',
            'core_product':
                'Geospatial analytics platform using satellite data + AI for infrastructure '
                'asset management: water networks, forestry, land stability, '
                'pipeline integrity. Products focus on condition monitoring and risk prediction.',
            'target_customers': 'Water utilities, forestry companies, infrastructure operators',
            'funding_revenue': 'Total raised: ~$21.1M. Majority stake acquired Dec 2024 (undisclosed acquirer).',
            'differentiators': [
                'Deep water utility vertical — network leakage and ground movement detection',
                'Multi-asset infrastructure (not just power lines)',
                'UK/EU regulatory alignment',
            ],
            'limitations': [
                'Narrow sector coverage (water, forestry)',
                'Small; post-acquisition strategic direction unclear',
                'No ESG or financial analytics',
            ],
            'relevance': 'Medium',
        },
    ]

    for c in companies:
        company_card(doc, c)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# CATEGORY 5: AI-NATIVE STARTUPS
# ─────────────────────────────────────────────────────────────────────────────
def add_cat5(doc):
    heading(doc, 'Category 5: AI-Native Startups — Physical Asset Intelligence & Geospatial AI for Finance/ESG', level=1)
    section_intro(doc,
        'The most strategically relevant category for the competitive analysis. These '
        'companies are building AI-native platforms that intersect geospatial/satellite '
        'data with financial, ESG, or industrial intelligence. They represent both '
        'competitive threats and potential partners or acquisition targets.')

    companies = [
        {
            'name': 'Kayrros',
            'founded': '2016',
            'hq': 'Paris, France (acquired Energy Aspects, 2026)',
            'stage': 'Acquired',
            'core_product':
                'MOST RELEVANT BENCHMARK. Satellite AI for energy asset-level emissions '
                'monitoring, oil/gas storage tracking, methane detection. Used for ESG '
                'compliance and commodity finance. See full profile in Category 1.',
            'target_customers': 'Energy firms, ESG investors, regulators, traders',
            'funding_revenue': '~$78M raised, $20.2M revenue (2024). Acquired 2026.',
            'differentiators': [
                'Only company that has successfully productised satellite AI into ESG/finance for energy assets',
                'Real-time monitoring + regulatory compliance use case',
                'Independent (non-company-reported) emissions data',
            ],
            'limitations': [
                'Energy sector only — no cross-sector asset intelligence',
                'Post-acquisition may become captive to Energy Aspects client base',
            ],
            'relevance': 'High',
        },
        {
            'name': 'MSCI GeoSpatial Asset Intelligence',
            'founded': '~2023 product launch within MSCI',
            'hq': 'New York, NY (MSCI division)',
            'stage': 'Product within Public Company (MSCI)',
            'core_product':
                'Location-based physical risk data for ~70,000 public and private companies. '
                'Maps company facility locations to nature risks and climate hazards. '
                'Enables investors and lenders to explore location-specific exposure. '
                'Recently launched dedicated GeoSpatial AI division.',
            'target_customers': 'Banks (lending risk), asset managers (portfolio), corporates (CSRD)',
            'funding_revenue': 'Part of MSCI ($2.5B total revenue). ESG analytics +40% YoY 2024. '
                               'Dedicated team hiring actively (Careers page confirms division)',
            'differentiators': [
                'Largest financial-institution customer base for physical asset risk',
                'Crosses into 70K private companies — beyond public markets',
                '70,000 companies with location-level data is unmatched at this scale',
                'Integration into existing portfolio and risk workflows',
            ],
            'limitations': [
                'Location data is modelled/inferred, not satellite-verified',
                'Hazard assessment is scenario-based — not operational monitoring',
                'Does not identify unknown or undisclosed assets',
                'Large company inertia — hard to iterate quickly vs nimble startups',
            ],
            'relevance': 'High',
        },
        {
            'name': 'Blackshark.ai',
            'founded': '2020',
            'hq': 'Vienna, Austria',
            'stage': 'Growth stage',
            'core_product':
                'Geospatial AI that extracts structured information from satellite/aerial '
                'imagery at global scale using ML. Builds a semantic 3D map of the world. '
                'Partners include Microsoft Flight Simulator; expanding to '
                'enterprise/industrial mapping.',
            'target_customers': 'Governments, infrastructure, gaming/simulation, enterprise',
            'funding_revenue': 'Funding details not publicly disclosed.',
            'differentiators': [
                'Global semantic extraction from satellite imagery — identifies objects/structures',
                'Technical depth in geospatial AI',
            ],
            'limitations': [
                'Not focused on financial/ESG use cases',
                'No monitoring/time-series asset intelligence',
            ],
            'relevance': 'Low',
        },
        {
            'name': 'Cervest',
            'founded': '2015',
            'hq': 'London, UK',
            'stage': 'Early-growth / restructured',
            'core_product':
                'Climate intelligence platform (EarthScan) providing AI-powered climate '
                'risk assessments for individual physical assets. Covers agriculture, '
                'infrastructure, commercial real estate. Mission: "Climate Security for global assets."',
            'target_customers': 'Corporates, insurers, financial institutions, governments',
            'funding_revenue': 'Raised ~$30M. Company underwent restructuring; current status uncertain.',
            'differentiators': [
                'Asset-level (not portfolio-level) climate risk — granular approach',
                'Covers multiple asset types: agriculture, infrastructure, real estate',
                'Strong academic/scientific backing (climate models)',
            ],
            'limitations': [
                'Company has faced funding challenges and restructuring',
                'Limited real-time satellite monitoring — primarily modelled risk',
                'Small team and limited commercial traction publicly visible',
            ],
            'relevance': 'Medium',
        },
        {
            'name': 'Orbital Sidekick',
            'founded': '2016',
            'hq': 'San Francisco, CA',
            'stage': 'Growth stage',
            'core_product':
                'Hyperspectral satellite analytics for oil & gas, mining, and industrial '
                'asset monitoring. GHOSt constellation provides molecular-level detection '
                'of hydrocarbons, methane, minerals from space. Key use cases: pipeline '
                'leak detection, mine tailings, spill monitoring.',
            'target_customers': 'Oil & gas operators, mining companies, environmental regulators',
            'funding_revenue': 'Raised ~$20-25M. Strategic investors include SK Telecom.',
            'differentiators': [
                'Hyperspectral = chemical/molecular detection from space — unique capability',
                'Identifies specific substances (methane, hydrocarbons, minerals) vs optical/SAR',
                'Direct ESG compliance use case for industrial operators',
            ],
            'limitations': [
                'Very small constellation — limited revisit frequency',
                'Narrow sector coverage (oil/gas, mining)',
                'Not yet productised for financial analytics or ESG ratings',
            ],
            'relevance': 'Medium',
        },
        {
            'name': 'Plume (YC W25)',
            'founded': '2024',
            'hq': 'Paris, France / San Francisco, CA',
            'stage': 'Seed (YC W25)',
            'core_product':
                'AI-powered geospatial platform for energy infrastructure site selection. '
                'Aggregates 150+ geographic datasets and AI agents to analyse documents, '
                'speed up permit processing and grid connection applications for renewable energy projects.',
            'target_customers': 'Renewable energy developers, infrastructure investors',
            'funding_revenue': 'Raised €3.3M (April 2026). Y Combinator W25 cohort. '
                               'Founded by ex-Palantir and ex-Harvard AI/geospatial researchers.',
            'differentiators': [
                'AI agents for unstructured regulatory/permit documents — novel approach',
                '150+ live geographic datasets integrated',
                'YC-backed with Palantir geospatial pedigree',
            ],
            'limitations': [
                'Very early stage — site selection only, not ongoing monitoring',
                'Renewable energy niche — not broad physical asset intelligence',
                'No revenue or customer traction publicly disclosed',
            ],
            'relevance': 'Medium',
        },
        {
            'name': 'RS Metrics',
            'founded': '2011',
            'hq': 'Bethesda, MD',
            'stage': 'Private / bootstrapped',
            'core_product':
                'Satellite-derived alternative data for financial markets. Products: '
                'MetalSignals (aluminium/steel production from space), RetailSignals (parking lot '
                'activity / foot traffic), Commodities Intelligence. Used by hedge funds and '
                'commodity traders for alpha generation.',
            'target_customers': 'Hedge funds, quant investors, commodity traders',
            'funding_revenue': 'Revenue and funding not publicly disclosed. Bootstrapped/profitable niche.',
            'differentiators': [
                'Proven product-market fit with quantitative hedge funds',
                'MetalSignals — industrial production intelligence from satellite',
                'Long track record (14+ years) — data history for backtesting',
            ],
            'limitations': [
                'Very niche — financial alpha, not ESG or operational intelligence',
                'No ESG, regulatory, or risk management products',
                'Limited geographic coverage (US-centric)',
                'Not designed for large-scale or cross-sector asset intelligence',
            ],
            'relevance': 'Medium',
        },
    ]

    for c in companies:
        company_card(doc, c)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# COMPETITIVE MATRIX
# ─────────────────────────────────────────────────────────────────────────────
def add_matrix(doc):
    heading(doc, 'Competitive Positioning Matrix', level=1)
    section_intro(doc,
        'The matrix below compares key players across six dimensions critical to '
        'the GeoAsset Intelligence value proposition. Scores: H = High capability / strength, '
        'M = Medium, L = Low / absent.')

    columns = [
        'Company',
        'Satellite\nMonitoring',
        'AI Analytics',
        'Multi-Sector\nAsset Coverage',
        'ESG/Finance\nIntegration',
        'Real-Time\n/ Near-RT',
        'Unknown\nAsset Discovery',
    ]

    rows_data = [
        # Company, sat, ai, multi, esg, rt, discovery
        ['Orbital Insight',          'H', 'H', 'H', 'M', 'M', 'H'],
        ['Kayrros',                  'H', 'H', 'L', 'H', 'H', 'M'],
        ['Planet Labs',              'H', 'M', 'H', 'L', 'M', 'M'],
        ['ICEYE',                    'H', 'M', 'M', 'L', 'H', 'M'],
        ['LiveEO',                   'H', 'H', 'L', 'L', 'M', 'L'],
        ['AiDash',                   'H', 'H', 'L', 'L', 'M', 'L'],
        ['MSCI GeoSpatial AI',       'L', 'M', 'H', 'H', 'L', 'L'],
        ['S&P Trucost (Physical)',   'L', 'M', 'H', 'H', 'L', 'L'],
        ['Clarity AI',               'L', 'H', 'M', 'H', 'L', 'L'],
        ['RS Metrics',               'H', 'M', 'L', 'L', 'M', 'M'],
        ['Satelytics',               'H', 'H', 'M', 'L', 'M', 'L'],
        ['Cervest',                  'L', 'M', 'M', 'M', 'L', 'L'],
        ['Orbital Sidekick',         'H', 'H', 'L', 'M', 'M', 'M'],
        ['TARGET STARTUP',           'H', 'H', 'H', 'H', 'H', 'H'],
    ]

    tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(columns))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'

    # Header row
    for j, col_label in enumerate(columns):
        cell = tbl.rows[0].cells[j]
        set_cell_bg(cell, DARK_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, col_label, bold=True, size=8, color=WHITE)

    # Data rows
    for i, row_data in enumerate(rows_data):
        row = tbl.rows[i + 1]
        is_target = row_data[0] == 'TARGET STARTUP'
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            if is_target:
                set_cell_bg(cell, ACCENT_BLUE)
            elif i % 2 == 0:
                set_cell_bg(cell, LIGHT_GREY)
            else:
                set_cell_bg(cell, WHITE)

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT

            if j == 0:
                add_run(p, val, bold=is_target, size=8.5,
                        color=WHITE if is_target else DARK_BLUE)
            else:
                color = GREEN if val == 'H' else (ORANGE if val == 'M' else RED_SOFT)
                if is_target:
                    color = WHITE
                add_run(p, val, bold=is_target, size=9, color=color)

    spacer(doc, 10)

    # Legend
    heading(doc, 'Matrix Legend', level=3)
    bullet(doc, 'H (High): Core capability, commercially deployed product, strong customer evidence')
    bullet(doc, 'M (Medium): Partial capability, early-stage product, or indirect coverage')
    bullet(doc, 'L (Low): Absent, modelled only, or not part of core product')
    bullet(doc, 'TARGET STARTUP: Aspirational positioning — combines all six dimensions in one platform',
           bold_prefix='Target Startup')

    spacer(doc, 8)
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# WHITE SPACE / OPPORTUNITY ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
def add_whitespace(doc):
    heading(doc, 'White Space & Strategic Opportunity Analysis', level=1)
    section_intro(doc,
        'Based on the competitive landscape, the following gaps represent addressable '
        'opportunities for a new entrant building at the intersection of geospatial AI, '
        'physical productive assets, and ESG/financial analytics.')

    gaps = [
        (
            '1. No cross-sector asset discovery platform',
            'Every satellite AI company focuses on a specific sector (energy, utilities, '
            'agriculture, infrastructure). No company currently offers automated discovery '
            'and classification of all types of productive physical assets (factories, '
            'mines, farms, facilities, power plants) from satellite imagery at global scale '
            'with a unified data model.',
            'An asset intelligence layer that identifies unknown/undisclosed assets and '
            'links them to ownership structures (via BvD/D&B type UBO data) represents '
            'a category-creating opportunity — especially for supply-chain due diligence, '
            'sanctions compliance, and CSRD reporting.',
        ),
        (
            '2. ESG data is disconnected from physical reality',
            'MSCI, Sustainalytics, and Clarity AI all rate companies based on self-reported '
            'data, regulatory filings, and modelled estimates. Kayrros is the only company '
            'that successfully bridges satellite monitoring → ESG-relevant metrics, but only '
            'for energy sector emissions.',
            'A platform that provides independently verified, satellite-derived ESG indicators '
            '(emissions, land use, water consumption, operational activity) for physical assets '
            'across all productive sectors — and maps them to financial instruments — '
            'addresses a USD multi-billion demand gap driven by CSRD, SFDR, and SEC climate rules.',
        ),
        (
            '3. Physical risk is modelled, not monitored',
            'S&P Trucost and MSCI GeoSpatial AI score 2.9M and 70K assets respectively, '
            'but both use climate scenario models, not actual real-time satellite data. '
            'This means a factory in a flood zone shows the same risk score whether it '
            'flooded last month or not.',
            'Real-time satellite monitoring combined with physical risk models enables '
            '"condition-adjusted" risk scoring — a step-change in accuracy for insurance '
            'underwriting, credit risk, and infrastructure maintenance.',
        ),
        (
            '4. Financial intelligence gap for private/unlisted productive assets',
            'Public equity markets are well covered. Private companies, unlisted '
            'infrastructure, farmland, and industrial facilities owned by private entities '
            'are largely invisible to financial data providers. Yet they represent the '
            'majority of global productive assets.',
            'Satellite-derived activity intelligence (production signals, utilisation rates, '
            'workforce proxies) for private assets creates differentiated alternative data '
            'for private equity, project finance, and ESG due diligence.',
        ),
        (
            '5. Supply-chain and sanctions enforcement gap',
            'Russia/Ukraine conflict, EUDR (EU Deforestation Regulation), Uyghur Forced '
            'Labor Prevention Act, and new EU supply-chain due diligence law (CSDDD) create '
            'immediate demand for physical verification of where goods are produced. '
            'Self-reported supply-chain data cannot be trusted.',
            'Satellite monitoring of specific industrial facilities linked to supply-chain '
            'graphs (who supplies whom) is a high-urgency compliance product with '
            'government-level and corporate-level demand. No pure-play product exists today.',
        ),
    ]

    for title, problem, opportunity in gaps:
        heading(doc, title, level=2)
        body(doc, 'PROBLEM: ' + problem)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(8)
        add_run(p, 'OPPORTUNITY: ', bold=True, size=10, color=GREEN)
        add_run(p, opportunity, size=10, color=BLACK)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# MARKET SIZING
# ─────────────────────────────────────────────────────────────────────────────
def add_market_sizing(doc):
    heading(doc, 'Market Sizing Reference Points', level=1)
    section_intro(doc,
        'Key market size data points to anchor the competitive analysis and investment thesis.')

    data = [
        ('Geospatial AI Market (2024)', 'USD 38B', 'USD 64.6B by 2030 (CAGR ~9%)', 'Precedence Research'),
        ('Satellite-based EO Market (2024)', 'USD 3.7B', 'USD 5.1B by 2030', 'Mordor Intelligence'),
        ('ESG Rating Services Market', '~$1.5B', 'Moderately concentrated; top 5 = ~65% of revenue', 'Mordor Intelligence'),
        ('S&P Global Market Intelligence', 'USD 4.65B (2024)', 'USD 4.92B (2025, +6%)', 'S&P Global Annual Report'),
        ('MSCI Total Revenue', '~USD 2.5B (2024)', 'ESG analytics +40% YoY 2024', 'MSCI Annual Report'),
        ('CoStar Group', 'USD 2.74B (2024)', 'USD 3.2B (2025, +19%)', 'CoStar Group IR'),
        ('Dun & Bradstreet', 'USD 2.38B (2024)', 'Being acquired by Clearlake for $7.7B', 'D&B IR'),
        ('Planet Labs Revenue', 'USD 244M (FY2024)', 'USD 281–289M guidance (FY2026)', 'Planet Labs IR'),
        ('Morningstar (incl. Sustainalytics)', '~USD 2.0–2.2B (2024)', 'Sustainalytics -6.5% YoY Q1 2025', 'Morningstar IR'),
        ('Palantir Revenue', 'USD 2.87B (FY2024)', 'USD 4.4B guidance (FY2025, +53%)', 'Palantir IR'),
    ]

    tbl = doc.add_table(rows=len(data) + 1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'

    headers = ['Company / Market', 'Current Scale', 'Trajectory', 'Source']
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        set_cell_bg(cell, DARK_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=9, color=WHITE)

    for i, row_data in enumerate(data):
        row = tbl.rows[i + 1]
        bg = LIGHT_GREY if i % 2 == 0 else WHITE
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            add_run(p, val, size=8.5, color=DARK_BLUE if j == 0 else BLACK,
                    bold=(j == 0))

    spacer(doc, 10)
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# SLIDE SCRIPT
# ─────────────────────────────────────────────────────────────────────────────
def add_slide_script(doc):
    heading(doc, 'Slide Deck Script — Competitive Analysis Presentation', level=1)
    section_intro(doc,
        'The following is a recommended slide-by-slide structure and speaker notes for '
        'presenting this competitive analysis to investors, advisors, or team members. '
        'Estimated deck length: 14 slides. Estimated presentation time: 20-25 minutes.')

    slides = [
        {
            'num': 1,
            'title': 'Title Slide',
            'visual': 'Company logo, tagline, presenter name, date. Background: satellite image of industrial landscape.',
            'content': [
                'Title: "The Physical World Has a Data Problem — We\'re Solving It"',
                'Subtitle: Competitive Landscape & Market Opportunity',
            ],
            'speaker_notes':
                'Open with the provocation: trillions of dollars of financial decisions are made '
                'about physical assets — factories, mines, farms, power plants — using data that '
                'is self-reported, static, and unverifiable. Our platform changes that.',
        },
        {
            'num': 2,
            'title': 'The Core Problem',
            'visual': '3-panel diagram: (1) Physical asset exists → (2) No real-time data → (3) Financial/ESG decisions made blindly.',
            'content': [
                '2.9 million assets mapped by S&P Trucost — but using modelled, not monitored data',
                'ESG ratings based 80%+ on self-reported company disclosures',
                'Unknown/undisclosed assets invisible to any financial data provider',
                'Supply chain due diligence laws (CSRD, CSDDD, EUDR, UFLPA) demand physical proof',
            ],
            'speaker_notes':
                'Frame the problem clearly: existing platforms know companies exist, but they '
                'do not watch assets. S&P Trucost has the largest physical risk dataset — 2.9M '
                'assets — but their data is modelled from climate scenarios, not observed by '
                'satellite. MSCI\'s GeoSpatial Asset Intelligence covers 70,000 companies — '
                'better, but still inferred location data. Nobody is watching. We will.',
        },
        {
            'num': 3,
            'title': 'Market Size & Tailwinds',
            'visual': 'Two bar charts: (1) Geospatial AI market growth $38B→$65B. (2) ESG data market growth. Plus 3 regulatory icons.',
            'content': [
                'Geospatial AI: USD 38B (2024) → USD 64.6B (2030), CAGR 9%',
                'Satellite EO: USD 3.7B (2024) → USD 5.1B (2030)',
                'ESG data market: ~$1.5B, top 5 players control 65%',
                'Regulatory demand: CSRD (50,000 EU companies), SFDR, EUDR, UFLPA, CSDDD',
            ],
            'speaker_notes':
                'This is a multi-billion-dollar market with strong regulatory tailwinds. '
                'Every EU company above 250 employees will report under CSRD by 2026. '
                'The demand is not speculative — it is legally mandated. '
                'No incumbent fully addresses the physical asset monitoring layer.',
        },
        {
            'num': 4,
            'title': 'Competitive Landscape Overview',
            'visual': '2x2 matrix: X-axis = Physical Asset Intelligence (low→high), Y-axis = ESG/Financial Integration (low→high). '
                     'Plot key players in quadrants.',
            'content': [
                'TOP RIGHT (target quadrant): Currently EMPTY — this is our position',
                'TOP LEFT: MSCI, Clarity AI, S&P Trucost — ESG but not physical monitoring',
                'BOTTOM RIGHT: Kayrros, AiDash, LiveEO — physical monitoring but narrow sector, weak ESG/finance link',
                'BOTTOM LEFT: SpaceKnow, RS Metrics — niche, limited integration',
            ],
            'speaker_notes':
                'This is the key insight slide. The top-right quadrant — high physical asset '
                'intelligence AND high ESG/financial integration — is empty. Kayrros came '
                'closest for the energy sector and was acquired for $X (undisclosed) by '
                'Energy Aspects in March 2026 after $78M in funding and $20M revenue. '
                'That exit validates the thesis. We are building the cross-sector version.',
        },
        {
            'num': 5,
            'title': 'Category 1: Satellite AI — Key Players',
            'visual': 'Logo grid of 6 companies with one-line descriptor and funding/revenue figure. '
                     'Color-code by relevance to our space.',
            'content': [
                'Orbital Insight: $132M raised, $44M revenue, acquired Privateer 2024 — pioneer, uncertain future',
                'Kayrros: $78M raised, $20M revenue, acquired 2026 — CLOSEST ANALOGUE',
                'Planet Labs: $244M revenue (FY2024) — data infrastructure, no analytics',
                'ICEYE: USD 2.8B valuation (2025), $430M+ raised — SAR, insurance/defence focus',
                'LiveEO: €72M+ raised, 3x revenue growth — utilities/infrastructure, Europe',
                'AiDash: $91.5M raised (Series C 2024) — utilities vegetation management',
            ],
            'speaker_notes':
                'The satellite AI space is consolidating. Three major acquisitions in 2024: '
                'Orbital Insight (by Privateer), Descartes Labs (by EarthDaily), Geosite '
                '(by Descartes Labs). Rezatec sold majority stake Dec 2024. This consolidation '
                'validates the value of these assets and reduces competition — acquirers are '
                'integrating specific verticals, not building the cross-sector platform.',
        },
        {
            'num': 6,
            'title': 'Category 2: Data Aggregators — Limitations',
            'visual': 'Table: Company | Asset Coverage | Real-time? | Satellite data? | Cross-sector?',
            'content': [
                'S&P Global Market Intelligence ($4.92B revenue 2025): Trucost maps 2.9M assets — modelled, not monitored',
                'MSCI ($2.5B revenue): GeoSpatial AI for 70K companies — inferred location, no satellite',
                'Dun & Bradstreet ($2.38B revenue 2024): 500M company records — zero asset monitoring',
                'Bureau van Dijk (Moody\'s): Best UBO/ownership data — no physical intelligence',
            ],
            'speaker_notes':
                'The data giants have distribution and customer relationships but are fundamentally '
                'limited: they track companies, not assets. They model risk, they do not monitor it. '
                'Their clients — the same banks and asset managers paying $100K+/year for data '
                'subscriptions — are already asking for this. The need to upgrade from company-level '
                'to asset-level, from modelled to monitored, is the commercial window.',
        },
        {
            'num': 7,
            'title': 'Category 3: ESG Platforms — The Data Gap',
            'visual': 'Funnel diagram showing: Raw satellite data → (GAP) → ESG Metric → Financial Decision. '
                     'Show where Clarity AI, MSCI ESG, Sustainalytics live.',
            'content': [
                'Clarity AI: $100M+ raised, BlackRock-backed, Forrester leader Q3 2024 — AI-native, but no satellite',
                'Sustainalytics (Morningstar): Revenue -6.5% YoY Q1 2025 — declining, legacy approach',
                'MSCI ESG: +40% YoY 2024 — winning by embedding geospatial, but not satellite-verified',
                'S&P Trucost: Physical risk for 99% of global market cap — but modelled only',
                'KEY GAP: No ESG platform uses real-time satellite monitoring of actual assets',
            ],
            'speaker_notes':
                'ESG data is a market in transition. Sustainalytics is declining. MSCI is winning '
                'by adding geospatial elements. Clarity AI is the sharpest AI-native challenger. '
                'But ALL of them have the same fundamental weakness: they rely on what companies '
                'say, not what satellites see. Regulators and auditors are increasingly challenging '
                'this. The EU CSRD auditing requirements will create demand for third-party physical '
                'verification — our product.',
        },
        {
            'num': 8,
            'title': 'Consolidation Wave — 2024 Acquisitions Signal Validation',
            'visual': 'Timeline of 2024-2026 acquisitions in the space with deal context.',
            'content': [
                'Apr 2024: Descartes Labs acquires Geosite (insurance geospatial)',
                'May 2024: Privateer acquires Orbital Insight ($132M raised, $44M ARR)',
                'Oct 2024: EarthDaily acquires Descartes Labs ($58M raised)',
                'Dec 2024: Unknown acquirer takes majority stake in Rezatec ($21M raised)',
                'Mar 2026: Energy Aspects acquires Kayrros ($78M raised, $20M ARR)',
                'INSIGHT: 5 major consolidations in 18 months — space is maturing rapidly',
            ],
            'speaker_notes':
                'This consolidation wave tells us two things: (1) the technology works — '
                'established companies are paying real money to acquire it. (2) The acquirers '
                'are vertical-specific — Privateer wants space sustainability, Energy Aspects '
                'wants energy intelligence. No one is building the horizontal platform. '
                'That is our opportunity and our differentiation.',
        },
        {
            'num': 9,
            'title': 'Our Differentiated Position — The 6-Dimension Stack',
            'visual': 'Radar/spider chart comparing TARGET STARTUP vs Kayrros vs MSCI GeoSpatial AI vs S&P Trucost '
                     'across: Satellite Monitoring / AI Analytics / Multi-Sector / ESG-Finance / Real-Time / Unknown Asset Discovery.',
            'content': [
                'Satellite Monitoring: Real-time, multi-source (optical + SAR + hyperspectral)',
                'AI Analytics: Asset identification, classification, activity monitoring, anomaly detection',
                'Multi-Sector: Industrial, energy, mining, agriculture, real estate — unified model',
                'ESG/Finance Integration: Output as ESG indicators, risk scores, alternative data',
                'Real-Time / Near-RT: Continuous monitoring, not annual snapshots',
                'Unknown Asset Discovery: Identify assets not in any existing database',
            ],
            'speaker_notes':
                'The spider chart will visually demonstrate that no existing player covers all '
                'six dimensions. Kayrros (our closest analogue) is strong on satellite + AI + '
                'ESG/finance + real-time, but only for energy. MSCI is strong on multi-sector + '
                'ESG/finance, but has no satellite monitoring and cannot discover unknown assets. '
                'We are the only platform that scores "High" on all six. This is the '
                'defensible white space.',
        },
        {
            'num': 10,
            'title': 'Target Customer Segments',
            'visual': '3-column layout: Segment | Pain today | Our solution | Willingness to pay reference.',
            'content': [
                'Financial institutions (banks, asset managers): Physical risk in portfolios is modelled, not verified → Real-time asset monitoring feed',
                'ESG data providers / rating agencies: Supplemental satellite data to verify self-reported disclosures → API integration',
                'Supply chain compliance teams: Cannot verify where tier-2/3 suppliers actually produce → Facility identification + monitoring',
                'Commodity traders / hedge funds: Alpha from industrial production signals → Alternative data subscription',
                'Regulators / government: Enforce EUDR, methane regulation, sanctions → Monitoring platform license',
            ],
            'speaker_notes':
                'We have five distinct customer segments, each with a different buying '
                'motion and value driver. Financial institutions are the highest-value '
                'segment (S&P Trucost, MSCI, Sustainalytics all charge six-figure '
                'subscriptions). Supply chain compliance is the fastest-growing use case '
                'driven by regulation. Commodity/hedge fund is the most immediate '
                'monetisable segment (RS Metrics proves the model at $6M+ revenue '
                'bootstrapped with a 30-person team).',
        },
        {
            'num': 11,
            'title': 'Competitive Moat — How We Stay Ahead',
            'visual': 'Flywheel diagram: More assets monitored → Better AI models → More customers → More data → ...',
            'content': [
                'Data moat: Every new asset monitored improves training data for AI models',
                'Network effects: Cross-sector linkage means we see supply-chain connections others miss',
                'Regulatory moat: First to build audit-grade satellite-verified ESG data creates switching costs',
                'Proprietary asset registry: Database of physical productive assets (undisclosed, private, unlisted) has no substitute',
                'Integration lock-in: Embed into MSCI/Bloomberg/Refinitiv workflows = distribution + retention',
            ],
            'speaker_notes':
                'The key moat is the proprietary asset registry. Once we have identified and '
                'classified 10 million productive assets globally — including assets that appear '
                'in no other database — that is a dataset with no substitute. Rebuilding it '
                'takes years and millions of dollars. Every customer who integrates our '
                'data into their risk models or compliance workflows creates switching costs.',
        },
        {
            'num': 12,
            'title': 'Go-to-Market Strategy',
            'visual': 'Staged GTM map: Phase 1 (beachhead vertical) → Phase 2 (sector expansion) → Phase 3 (platform).',
            'content': [
                'Phase 1 — Beachhead: Energy/industrial emissions monitoring (regulatory compliance driver, PROVEN by Kayrros exit)',
                'Phase 2 — Sector expansion: Mining, agriculture, manufacturing (add sectors to unified asset model)',
                'Phase 3 — Platform: Become the physical asset intelligence layer for ESG data providers and financial institutions',
                'Distribution: Partner with MSCI, Bloomberg, Refinitiv as data distribution channels (vs direct sales only)',
                'Regulatory entry: EU agencies, EUDR enforcement as first government customers',
            ],
            'speaker_notes':
                'We start with energy/industrial emissions because: (1) the market is proven '
                '— Kayrros built $20M ARR here before acquisition. (2) EU Methane Regulation '
                'creates immediate compliance demand. (3) The data infrastructure (satellites, '
                'AI models) needed for energy applies directly to other sectors. From there '
                'we expand sector by sector, building the most comprehensive physical asset '
                'register ever created.',
        },
        {
            'num': 13,
            'title': 'Key Risks & Mitigations',
            'visual': 'Risk matrix (likelihood vs impact) with 5 risks plotted.',
            'content': [
                'MSCI/S&P build it themselves → Mitigate: Move fast, build asset registry moat, they are slow to iterate',
                'Satellite data access costs → Mitigate: Partner with Planet/ICEYE rather than own constellation',
                'AI model accuracy at scale → Mitigate: Human-in-the-loop validation, gradual sector rollout',
                'Regulatory uncertainty on ESG mandates → Mitigate: Diversify across regulatory frameworks (EU + US + UK)',
                'Customer data privacy / sovereignty → Mitigate: Focus on publicly observable physical assets, transparent methodology',
            ],
            'speaker_notes':
                'The incumbent risk is real — MSCI has already launched "GeoSpatial Asset '
                'Intelligence" with 70K companies and is hiring dedicated teams. But they are '
                'building on top of their existing methodology (company-reported + inferred) '
                'rather than satellite-first. Their moat is distribution; our moat must be '
                'the quality and independence of the underlying physical observation. '
                'We need to move faster than they can evolve.',
        },
        {
            'num': 14,
            'title': 'Summary: The Opportunity',
            'visual': 'Single clean summary graphic with 3 boxes: Problem → Gap → Our Solution.',
            'content': [
                'PROBLEM: $trillions in financial/ESG decisions made using self-reported, unverified physical asset data',
                'GAP: No platform combines satellite AI + multi-sector asset intelligence + ESG/finance integration at scale',
                'PROOF POINTS: 5 major acquisitions 2024-2026 validate the category; Kayrros exit is the sector-specific proof of concept',
                'OUR SOLUTION: The first cross-sector, satellite-verified, AI-powered physical asset intelligence platform',
                'CALL TO ACTION: [Investment ask / partnership proposal / pilot programme]',
            ],
            'speaker_notes':
                'Close with conviction: the market is real, the demand is regulatory-mandated, '
                'the incumbents have structural gaps, and the acquisition wave proves the exit '
                'paths. We are building the platform that the Kayrros acquirer, the MSCI '
                'GeoSpatial team, and the S&P Trucost team all wish they had. '
                'The question is not whether this market exists — it\'s whether we build '
                'the defining company in it.',
        },
    ]

    for slide in slides:
        heading(doc, f"Slide {slide['num']}: {slide['title']}", level=2)

        # Visual description
        vp = doc.add_paragraph()
        vp.paragraph_format.space_before = Pt(2)
        vp.paragraph_format.space_after = Pt(4)
        add_run(vp, 'VISUAL: ', bold=True, size=9.5, color=ACCENT_BLUE)
        add_run(vp, slide['visual'], italic=True, size=9.5, color=MID_GREY)

        # Content bullets
        heading(doc, 'Slide Content', level=3)
        for item in slide['content']:
            bullet(doc, item)

        # Speaker notes
        heading(doc, 'Speaker Notes', level=3)
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after = Pt(10)
        sp.paragraph_format.left_indent = Inches(0.25)
        add_run(sp, slide['speaker_notes'], size=9.5, color=BLACK, italic=True)
        spacer(doc, 6)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# SOURCES
# ─────────────────────────────────────────────────────────────────────────────
def add_sources(doc):
    heading(doc, 'Sources & References', level=1)
    section_intro(doc, 'All data points sourced May 2026. Financial figures reflect most recently available public disclosures.')

    sources = [
        ('Orbital Insight acquisition by Privateer', 'SpaceInsider, May 2024 — https://spaceinsider.tech/2024/05/06/privateer-raises-56-5-million-and-acquires-orbital-insight/'),
        ('Orbital Insight revenue ~$44M', 'Latka, June 2024 — https://getlatka.com/companies/orbital-insight'),
        ('Kayrros revenue $20.2M, 80 customers', 'Latka, October 2024 — https://getlatka.com/companies/kayrros'),
        ('Kayrros acquisition by Energy Aspects', 'Kayrros press release, March 2026 — https://www.kayrros.com/press-release/'),
        ('Kayrros Fortune Change the World #30', 'Kayrros, 2024 — https://www.kayrros.com/kayrros-rises-to-30-in-fortunes-change-the-world-2024-list/'),
        ('Planet Labs FY2024 revenue $244M; FY2026 guidance', 'Planet Labs IR / Finimize — https://finimize.com/content/pl-asset-snapshot'),
        ('Planet Labs $280M Germany contract', 'Various reports, July 2025'),
        ('ICEYE EUR 150M round, EUR 2.4B valuation', 'The Fast Mode, Dec 2025 — https://www.thefastmode.com/investments-and-expansions/46254-iceye-strengthens-sovereign-sar-satellite-capability-with-new-150m-funding'),
        ('ICEYE 54 SAR satellites', 'BeInsure, 2025'),
        ('Descartes Labs acquisition by EarthDaily', 'Via Satellite, October 2024 — https://www.satellitetoday.com/imagery-and-sensing/2024/10/15/earthdaily-acquires-descartes-labs/'),
        ('Satellogic FY2024 revenue $12.9M (+28% YoY)', 'Satellogic Press Release, 2025'),
        ('Satellogic $30M defence contract', 'Satellogic, April 2025 — https://satellogic.com/news/press-releases/'),
        ('SpaceKnow funding $5.45M (Series A, 2017)', 'CB Insights / Tracxn'),
        ('Satelytics revenue $6.1M (Oct 2024)', 'Latka — https://getlatka.com/companies/satelytics'),
        ('Ursa Space total raised $69.7M; Aug 2025 round', 'Tracxn / CB Insights'),
        ('LiveEO €25M Series B (Jun 2024) + €28M new round', 'TNW, VentureBurn, ESA — https://thenextweb.com/news/liveeo-28m-funding-civil-infrastructure-defence-twinspector'),
        ('AiDash $91.5M total raised (Series C 2024)', 'Dealroom'),
        ('Rezatec $21.1M raised; majority stake Dec 2024', 'PitchBook / Tracxn'),
        ('S&P Global FY2025 revenue $15.3B; Market Intelligence $4.92B', 'S&P Global Annual Report 2025 / MacroTrends'),
        ('S&P Trucost: 2.9M+ assets, 15K+ companies, 7 hazards', 'S&P Global / WRDS data announcement'),
        ('MSCI ESG +40% YoY 2024; Total revenue ~$2.5B', 'MSCI 2024 Annual Report'),
        ('MSCI GeoSpatial Asset Intelligence: 70K companies', 'MSCI Careers page / competitive landscape search'),
        ('Sustainalytics revenue -6.5% Q1 2025 to $28.8M', 'Morningstar Q1 2025 earnings'),
        ('Dun & Bradstreet FY2024 revenue $2.38B; Clearlake acquisition $7.7B', 'D&B IR, March 2025'),
        ('Clarity AI $100-120M raised; BlackRock investor; Forrester Wave leader Q3 2024', 'CB Insights / Tracxn / Clarity AI press releases'),
        ('Clarity AI acquired ecolytiq', 'ESG Today, July 2025'),
        ('CoStar Group FY2025 revenue $3.2B (+19%); FY2024 $2.74B', 'CoStar Group IR, 2025'),
        ('Geospatial AI market $38B (2024) → $64.6B (2030)', 'Precedence Research'),
        ('Satellite EO market $3.7B (2024) → $5.1B (2030)', 'Mordor Intelligence'),
        ('ESG rating services: top 5 = 65% market share', 'Mordor Intelligence'),
        ('Plume €3.3M seed, YC W25, ex-Palantir founders', 'EU-Startups, April 2026 / BeBeez'),
        ('Palantir FY2024 revenue $2.87B; FY2025 guidance $4.4B', 'Palantir IR, 2025'),
    ]

    tbl = doc.add_table(rows=len(sources) + 1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for j, h in enumerate(['Data Point', 'Source / URL']):
        cell = tbl.rows[0].cells[j]
        set_cell_bg(cell, DARK_BLUE)
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, size=9, color=WHITE)

    for i, (label, source) in enumerate(sources):
        row = tbl.rows[i + 1]
        bg = LIGHT_GREY if i % 2 == 0 else WHITE
        for j, val in enumerate([label, source]):
            cell = row.cells[j]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            add_run(p, val, size=8, color=DARK_BLUE if j == 0 else BLACK, bold=(j == 0))


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────
def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

    # Default paragraph font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    add_cover(doc)
    add_exec_summary(doc)
    add_cat1(doc)
    add_cat2(doc)
    add_cat3(doc)
    add_cat4(doc)
    add_cat5(doc)
    add_matrix(doc)
    add_whitespace(doc)
    add_market_sizing(doc)
    add_slide_script(doc)
    add_sources(doc)

    out_path = '/home/genmsadmin/Projects/geoasset-location/GeoAsset_Competitive_Analysis.docx'
    doc.save(out_path)
    print(f'Document saved to: {out_path}')
    return out_path


if __name__ == '__main__':
    main()
