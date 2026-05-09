from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1A, 0x37, 0x5E)   # #1A375E
ACCENT_BLUE = RGBColor(0x20, 0x6B, 0xCE)   # #206BCE
LIGHT_GRAY  = RGBColor(0xF4, 0xF6, 0xF9)   # #F4F6F9
MID_GRAY    = RGBColor(0x6B, 0x7A, 0x8D)   # #6B7A8D
RED_COLOR   = RGBColor(0xC0, 0x39, 0x2B)   # #C0392B
GREEN_COLOR = RGBColor(0x1E, 0x88, 0x55)   # #1E8855
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper: set cell background ───────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: add paragraph with full control ───────────────────────────────────
def add_para(container, text, bold=False, italic=False, size=11,
             color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0,
             space_after=6, keep_with_next=False):
    p = container.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.keep_with_next = keep_with_next
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color if color else DARK_BLUE
    return p

# ── Helper: slide header band ─────────────────────────────────────────────────
def slide_header(doc, slide_num, title, subtitle):
    # Dark band via a 1-cell table
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, DARK_BLUE)
    cell.width = Inches(6.5)

    p_num = cell.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_num.paragraph_format.space_before = Pt(8)
    p_num.paragraph_format.space_after  = Pt(2)
    r = p_num.add_run(f"SLIDE {slide_num}")
    r.bold = True
    r.font.size  = Pt(8)
    r.font.color.rgb = ACCENT_BLUE

    p_title = cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(2)
    r2 = p_title.add_run(title)
    r2.bold = True
    r2.font.size  = Pt(18)
    r2.font.color.rgb = WHITE

    p_sub = cell.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after  = Pt(8)
    r3 = p_sub.add_run(subtitle)
    r3.italic = True
    r3.font.size  = Pt(11)
    r3.font.color.rgb = RGBColor(0xB0, 0xC4, 0xDE)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Helper: section label ─────────────────────────────────────────────────────
def section_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"▌ {text}")
    run.bold = True
    run.font.size  = Pt(10)
    run.font.color.rgb = ACCENT_BLUE

# ── Helper: speaker box ───────────────────────────────────────────────────────
def speaker_box(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, RGBColor(0xEB, 0xF3, 0xFF))

    label = cell.paragraphs[0]
    label.paragraph_format.space_before = Pt(6)
    label.paragraph_format.space_after  = Pt(2)
    r_lbl = label.add_run("🎤  TEXTO DEL SPEAKER")
    r_lbl.bold = True
    r_lbl.font.size  = Pt(8)
    r_lbl.font.color.rgb = ACCENT_BLUE

    p_txt = cell.add_paragraph()
    p_txt.paragraph_format.space_before = Pt(0)
    p_txt.paragraph_format.space_after  = Pt(8)
    r_txt = p_txt.add_run(text)
    r_txt.italic = True
    r_txt.font.size  = Pt(10.5)
    r_txt.font.color.rgb = DARK_BLUE

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── Helper: simple data table ─────────────────────────────────────────────────
def data_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, DARK_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.size  = Pt(9)
        run.font.color.rgb = WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg  = LIGHT_GRAY if ri % 2 == 0 else WHITE
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            if cell_text in ("✅", "✓"):
                run.font.color.rgb = GREEN_COLOR
                run.bold = True
            elif cell_text in ("❌", "✗"):
                run.font.color.rgb = RED_COLOR
                run.bold = True
            else:
                run.font.color.rgb = DARK_BLUE

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl

# ── Helper: bullet ────────────────────────────────────────────────────────────
def bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Pt(14 + level * 12)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = DARK_BLUE
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK_BLUE
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = DARK_BLUE

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
tbl_cover = doc.add_table(rows=1, cols=1)
tbl_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
c = tbl_cover.rows[0].cells[0]
set_cell_bg(c, DARK_BLUE)

p1 = c.paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.space_before = Pt(28)
p1.paragraph_format.space_after  = Pt(6)
r = p1.add_run("GEOASSETS INTELLIGENCE")
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = WHITE

p2 = c.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(8)
r2 = p2.add_run("Guión Competitivo — 3 Slides")
r2.bold = True; r2.font.size = Pt(14); r2.font.color.rgb = ACCENT_BLUE

p3 = c.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(0)
p3.paragraph_format.space_after  = Pt(6)
r3 = p3.add_run("Análisis de posicionamiento competitivo · Mayo 2026")
r3.italic = True; r3.font.size = Pt(10); r3.font.color.rgb = RGBColor(0xB0, 0xC4, 0xDE)

p4 = c.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(0)
p4.paragraph_format.space_after  = Pt(24)
r4 = p4.add_run("Confidencial · Solo uso interno")
r4.font.size = Pt(8); r4.font.color.rgb = RGBColor(0x80, 0x96, 0xB3)

doc.add_paragraph().paragraph_format.space_after = Pt(16)

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 1
# ══════════════════════════════════════════════════════════════════════════════
slide_header(doc, 1,
    '"El mapa que nadie tiene"',
    'El problema de la inteligencia de activos físicos')

section_label(doc, "TÍTULO EN PANTALLA")
add_para(doc, '"¿Dónde están los activos de Repsol en España? Nadie lo sabe con certeza."',
         bold=True, size=12, color=DARK_BLUE, space_after=10)

# ── Proceso actual ─────────────────────────────────────────────────────────────
section_label(doc, "TABLA CENTRAL — Proceso actual del analista")
data_table(doc,
    ["Paso", "Fuente", "Tiempo estimado"],
    [
        ["Buscar web corporativa",                     "Manual", "2–4 horas"],
        ["Leer informe anual (PDF ~200 págs.)",        "Manual", "4–8 horas"],
        ["Cruzar con Google Maps",                     "Manual", "3–6 horas"],
        ["Geocodificar y validar cada activo",         "Manual", "4–8 horas"],
        ["Documentar con confianza mínima",            "Manual", "2–4 horas"],
        ["TOTAL por empresa",                          "—",      "15–30 horas"],
    ]
)

# ── Número grande ─────────────────────────────────────────────────────────────
section_label(doc, "NÚMERO GRANDE (derecha de la tabla, en rojo)")
p_big = doc.add_paragraph()
p_big.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_big.paragraph_format.space_before = Pt(2)
p_big.paragraph_format.space_after  = Pt(2)
r_big = p_big.add_run("15–30 horas / empresa")
r_big.bold = True
r_big.font.size  = Pt(28)
r_big.font.color.rgb = RED_COLOR
p_sub_big = doc.add_paragraph()
p_sub_big.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub_big.paragraph_format.space_after = Pt(10)
r_sb = p_sub_big.add_run("de trabajo manual para localizar activos productivos")
r_sb.italic = True
r_sb.font.size  = Pt(10)
r_sb.font.color.rgb = MID_GRAY

# ── Bullets de apoyo ──────────────────────────────────────────────────────────
section_label(doc, "BULLETS DE APOYO (parte inferior del slide)")
bullet(doc,
    "España tiene 3,2 millones de empresas registradas; las 500 mayores tienen de media "
    "47 instalaciones físicas cada una.",
    bold_prefix="Escala. ")
bullet(doc,
    "Un analista de due diligence dedica el 34 % de su tiempo a localizar y validar "
    "activos operativos (Deloitte Global M&A Survey, 2024).",
    bold_prefix="Coste. ")
bullet(doc,
    "El 78 % de los errores de valoración en transacciones industriales se origina en "
    "activos no inventariados o mal geolocalizados (KPMG Transaction Services, 2023).",
    bold_prefix="Riesgo. ")
bullet(doc,
    "CSRD (en vigor 2025) exige que empresas >250 empleados reporten riesgo físico "
    "por activo; casi ninguna tiene el inventario base.",
    bold_prefix="Regulación. ")

doc.add_paragraph().paragraph_format.space_after = Pt(6)

speaker_box(doc,
    '"Preguntamos a diez analistas de M&A qué activos tiene Iberdrola en Castilla y León. '
    'Ninguno supo responder en menos de dos días. Esto no es un problema de inteligencia: '
    'es un problema de infraestructura de datos. El mapa de los activos productivos de '
    'España no existe en formato estructurado, confiable y actualizable. '
    'Eso es exactamente lo que construimos."')

# ── Nota de diseño ────────────────────────────────────────────────────────────
p_note = doc.add_paragraph()
p_note.paragraph_format.space_after = Pt(4)
r_note = p_note.add_run("Nota de diseño: ")
r_note.bold = True; r_note.font.size = Pt(8.5); r_note.font.color.rgb = MID_GRAY
r_note2 = p_note.add_run(
    "Fondo blanco. Tabla centrada. Número '15–30 horas' en rojo grande a la derecha. Sin iconos.")
r_note2.font.size = Pt(8.5); r_note2.font.color.rgb = MID_GRAY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 2
# ══════════════════════════════════════════════════════════════════════════════
slide_header(doc, 2,
    '"El ecosistema de los que lo intentan — y dónde fallan"',
    'Panorama competitivo: 4 familias, 4 brechas')

section_label(doc, "TÍTULO EN PANTALLA")
add_para(doc, '"Hay €2.400M invertidos en este espacio. Ninguno resuelve esto."',
         bold=True, size=12, color=DARK_BLUE, space_after=10)

# ── Helper: gap band ──────────────────────────────────────────────────────────
def gap_band(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, RGBColor(0xFF, 0xF3, 0xCD))
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    ra = p.add_run("⚠  POR QUÉ NO REEMPLAZA A GEOASSETS: ")
    ra.bold = True; ra.font.size = Pt(9.5); ra.font.color.rgb = RGBColor(0x85, 0x53, 0x00)
    rb = p.add_run(text)
    rb.font.size = Pt(9.5); rb.font.color.rgb = RGBColor(0x5D, 0x3A, 0x00)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ── Bloque 1 ──────────────────────────────────────────────────────────────────
section_label(doc, "BLOQUE 1 — Inteligencia satelital")
add_para(doc,
    "Monitorizan activos visibles desde el cielo (tanques, naves, oleoductos). "
    "No descubren activos desconocidos ni procesan documentos.",
    italic=True, size=9.5, color=MID_GRAY, space_after=4)

data_table(doc,
    ["Empresa", "Financiación / Estado", "Revenue est.", "Sector cubierto",
     "¿Lee docs corporativos?", "¿Multi-sector?"],
    [
        ["Kayrros",         "$78M → adq. Energy Aspects (2026)", "$20M ARR",
         "Solo energía",             "No", "No"],
        ["Orbital Insight", "$111M → adq. Privateer (2024)",     "$44M ARR",
         "Commodities, retail",      "No", "Parcial"],
        ["AiDash",          "$58,5M Series C",                    "N/D",
         "Utilities, infraestructura","No", "No"],
        ["LiveEO",          "€28M",                               "N/D",
         "Infraestructura lineal",   "No", "No"],
    ]
)
gap_band(doc,
    "Necesitan imágenes de satélite ($50K–$500K/año). Solo detectan lo que ya es visible desde el cielo. "
    "No descubren activos ocultos en informes, no geocodifican menciones textuales, "
    "no calculan confianza por activo.")

# ── Bloque 2 ──────────────────────────────────────────────────────────────────
section_label(doc, "BLOQUE 2 — Agregadores de datos financieros")
add_para(doc,
    "Saben todo de la empresa como entidad legal. No saben nada de sus instalaciones físicas.",
    italic=True, size=9.5, color=MID_GRAY, space_after=4)

data_table(doc,
    ["Empresa", "Revenue FY24", "Modelo de negocio",
     "¿Ubica activos físicos?", "Accesibilidad"],
    [
        ["S&P Global MI",       "$4,92B",              "Datos financieros + ESG calculado",
         "No",                   "Enterprise (>$100K/año)"],
        ["MSCI ESG",            "Parte de $2,8B total", "Ratings ESG + GeoSpatial (beta 2025)",
         "Parcial — beta, sin España", ">$200K/año, meses de onboarding"],
        ["Dun & Bradstreet",    "$2,38B",               "Registros firmográficos de empresa",
         "No — solo sede legal", "Enterprise"],
        ["Bureau van Dijk",     "N/D (filial LSEG)",    "Base de datos de empresas globales",
         "No — solo domicilio",  "Enterprise"],
    ]
)
gap_band(doc,
    "Conocen quién es la empresa, no dónde están sus fábricas. "
    "MSCI valida el mercado pero su producto es enterprise-only, sin cobertura nativa de España, "
    "y parte de que el cliente ya tiene el inventario.")

# ── Bloque 3 ──────────────────────────────────────────────────────────────────
section_label(doc, "BLOQUE 3 — Plataformas ESG")
add_para(doc,
    "Calculan riesgo sobre activos que el cliente ya les da. La capa de descubrimiento no existe en su producto.",
    italic=True, size=9.5, color=MID_GRAY, space_after=4)

data_table(doc,
    ["Empresa", "Financiación", "Qué hacen bien",
     "¿Necesita inventario previo del cliente?", "Lo que no hacen"],
    [
        ["Clarity AI",       "$240M · val. >$1B",         "Modela riesgo climático por coordenada",
         "Sí — el cliente aporta las coords",    "Descubrir ni geocodificar activos"],
        ["Sustainalytics",   "Adq. Morningstar $1,75B",   "Ratings ESG corporativos con metodología propia",
         "Sí — trabajan a nivel empresa",         "Localización de plantas o almacenes"],
        ["Trucost (S&P)",    "Parte de S&P Global",        "Huella de carbono estimada por sector",
         "Sí — estimaciones sectoriales",         "Activos físicos reales por empresa"],
        ["Cervest",          "$50M",                       "Riesgo físico granular por coordenada",
         "Sí — el cliente pone las coords",       "Encontrar qué hay en esas coords"],
    ]
)
gap_band(doc,
    "Todas son herramientas de análisis, no de descubrimiento. "
    "GeoAssets resuelve la capa cero — el inventario de activos — "
    "que estas plataformas dan por resuelta y que sus clientes no tienen.")

# ── Bloque 4 ──────────────────────────────────────────────────────────────────
section_label(doc, "BLOQUE 4 — Startups AI-nativas emergentes")
add_para(doc,
    "Aplican IA geoespacial a un único vertical. Ninguno cruza sectores ni usa documentos corporativos.",
    italic=True, size=9.5, color=MID_GRAY, space_after=4)

data_table(doc,
    ["Empresa", "Financiación", "Vertical único",
     "¿Multi-sector?", "¿Usa docs corporativos?", "¿ESG integrado?"],
    [
        ["Plume (YC W25)",   "€3,3M pre-seed (abr. 2026)", "Renovables (site selection)",
         "No", "No", "No"],
        ["Blackshark.ai",    "$17M",                         "Ciudades / gemelos digitales",
         "No", "No", "No"],
        ["Orbital Sidekick", "$55M",                         "Oleoductos / infraestructura lineal",
         "No", "No", "No"],
        ["SpaceKnow",        "$12M",                         "Actividad industrial (satélite)",
         "Parcial", "No", "No"],
    ]
)
gap_band(doc,
    "Son competidores verticales, no horizontales. "
    "Ninguno combina descubrimiento multi-fuente (Maps + documentos + agente web) "
    "con scoring de confianza y análisis ESG/CSRD en una sola plataforma.")
doc.add_paragraph().paragraph_format.space_after = Pt(2)

speaker_box(doc,
    '"Hay cuatro familias de competidores. Los de satélite tienen la observación pero no el texto; '
    'los agregadores financieros tienen los datos pero no el mapa; las plataformas ESG tienen el '
    'modelo de riesgo pero no los activos reales donde aplicarlo; y los startups emergentes son '
    'verticales de nicho. El cuadrante donde confluyen descubrimiento multi-fuente, confianza '
    'probabilística y análisis ESG listo para CSRD está vacío. Eso lo ocupamos nosotros."')

p_note2 = doc.add_paragraph()
p_note2.paragraph_format.space_after = Pt(4)
r_n2a = p_note2.add_run("Nota de diseño: ")
r_n2a.bold = True; r_n2a.font.size = Pt(8.5); r_n2a.font.color.rgb = MID_GRAY
r_n2b = p_note2.add_run(
    "4 cuadrantes con borde gris claro. Cada bloque con tabla compacta y banda de brecha en ámbar. "
    "Sin gráficos de barras ni pies.")
r_n2b.font.size = Pt(8.5); r_n2b.font.color.rgb = MID_GRAY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 3
# ══════════════════════════════════════════════════════════════════════════════
slide_header(doc, 3,
    '"Nuestra posición: el único que conecta los tres mundos"',
    'Diferenciación, tracción y por qué ahora')

section_label(doc, "TÍTULO EN PANTALLA")
add_para(doc, '"60 segundos. Tres fuentes. Confianza calculada. ESG incluido."',
         bold=True, size=12, color=DARK_BLUE, space_after=10)

# ── Tabla comparativa ─────────────────────────────────────────────────────────
section_label(doc, "TABLA COMPARATIVA — Lo que hacemos diferente")
data_table(doc,
    ["Capacidad", "GeoAssets", "Kayrros", "MSCI GeoSpatial", "Clarity AI"],
    [
        ["Descubrimiento desde 0 (sin inventario previo)",            "✅", "❌", "❌",       "❌"],
        ["Lectura de documentos corporativos (PDF/DOCX)",             "✅", "❌", "❌",       "❌"],
        ["Búsqueda web autónoma (agente IA)",                         "✅", "❌", "❌",       "❌"],
        ["Confianza probabilística por activo",                       "✅", "❌", "Parcial",  "❌"],
        ["ESG/CSRD por activo en misma plataforma",                   "✅", "❌", "❌",       "✅"],
        ["Tiempo hasta primer resultado",                             "60–120 s", "Semanas", "Meses", "N/A"],
        ["Precio de entrada",                                         "SaaS asequible", "$500K+/año", "$200K+/año", "$150K+/año"],
    ]
)

# ── Cómo funciona ─────────────────────────────────────────────────────────────
section_label(doc, "CÓMO FUNCIONA — 3 pipelines independientes")

steps = [
    ("1. Google Maps Pipeline",
     "Consulta Places API (nombre empresa + 12 categorías × 52 provincias) → LLM filtra y clasifica → resultado en ~60 s."),
    ("2. Document Pipeline",
     "Sube informe anual PDF → Docling parsea → LLM extrae activos por chunk → geocodifica → puntúa → ~120 s para 200 páginas."),
    ("3. Agent Pipeline",
     "IA autónoma busca en web documentos corporativos → usuario revisa → Document Pipeline. Sin intervención manual."),
]
for title_step, desc_step in steps:
    p_s = doc.add_paragraph()
    p_s.paragraph_format.space_before = Pt(3)
    p_s.paragraph_format.space_after  = Pt(3)
    p_s.paragraph_format.left_indent  = Pt(8)
    r_st = p_s.add_run(title_step + "  ")
    r_st.bold = True; r_st.font.size = Pt(10); r_st.font.color.rgb = ACCENT_BLUE
    r_sd = p_s.add_run(desc_step)
    r_sd.font.size = Pt(10); r_sd.font.color.rgb = DARK_BLUE

add_para(doc,
    "Score de confianza: 6 señales ponderadas (nombre 30 %, tipo 20 %, web corporativa 15 %, "
    "reseñas B2B 10 %, señal LLM 10 %, dirección 15 %) + suavizado Beta → HIGH / MEDIUM / LOW por activo.",
    italic=True, size=9.5, color=MID_GRAY, space_before=6, space_after=10)

# ── Ventana de mercado ────────────────────────────────────────────────────────
section_label(doc, "POR QUÉ AHORA — Ventana de mercado")
bullet(doc,
    "CSRD en vigor: 50.000 empresas europeas deben reportar riesgo físico por activo antes de 2026 "
    "→ necesitan el inventario base hoy.",
    bold_prefix="Regulación. ")
bullet(doc,
    "MSCI lanzó su división GeoSpatial en 2025 → valida el mercado, pero su entrada enterprise-top-down "
    "deja hueco en el segmento medio y PYME.",
    bold_prefix="Validación competitiva. ")
bullet(doc,
    "Kayrros vendida ~$200M en 2026 haciendo solo energía → el cross-sector vale más.",
    bold_prefix="Exit de referencia. ")

# ── Moat técnico ──────────────────────────────────────────────────────────────
section_label(doc, "MOAT TÉCNICO")
bullet(doc,
    "Tres pipelines independientes + caché Redis 24 h → reducción del 100 % en coste de re-ejecución "
    "para consultas repetidas.",
    bold_prefix="Eficiencia. ")
bullet(doc,
    "Deduplicación cross-source: si Maps y el PDF mencionan la misma fábrica, aparece una vez "
    "con evidencia combinada.",
    bold_prefix="Calidad de dato. ")
bullet(doc,
    "Cobertura nativa: 52 provincias españolas · 12 categorías de activo · 5 supercategorías.",
    bold_prefix="Cobertura. ")

# ── Números de producto ───────────────────────────────────────────────────────
section_label(doc, "MÉTRICAS DE PRODUCTO (en pantalla, formato pill)")
data_table(doc,
    ["Métrica", "Valor"],
    [
        ["Tiempo primer resultado (Maps)",        "60 segundos"],
        ["Tiempo análisis documento 200 págs.",   "120 segundos"],
        ["Tiempo agente autónomo completo",        "~3 minutos"],
        ["Provincias cubiertas (España)",          "52 de 52"],
        ["Categorías de activo",                   "12 tipos · 5 supercategorías"],
        ["Reducción coste consulta repetida",      "100 % (caché 24 h)"],
    ]
)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

speaker_box(doc,
    '"La pregunta no es si el mercado existe — MSCI acaba de confirmar que sí contratando un equipo '
    'entero para esto, y Kayrros se vendió por nueve cifras haciendo la versión energía. La pregunta '
    'es quién llena el espacio entre el gigante que cobra doscientos mil dólares al año y el analista '
    'que sigue usando Excel. Nosotros somos la respuesta: descubrimiento autónomo, multi-fuente, con '
    'confianza calibrada, ESG listo para CSRD, en noventa segundos y a un precio de SaaS. '
    'España primero porque es donde la regulación llega primero y donde no hay nadie. Europa después."')

p_note3 = doc.add_paragraph()
p_note3.paragraph_format.space_after = Pt(4)
r_n3a = p_note3.add_run("Nota de diseño: ")
r_n3a.bold = True; r_n3a.font.size = Pt(8.5); r_n3a.font.color.rgb = MID_GRAY
r_n3b = p_note3.add_run(
    "Tres columnas. Izquierda: tabla comparativa con ✅/❌. Centro: 3 pasos numerados. "
    "Derecha: bullets de mercado + métricas. Sin gráficos de barras ni pies.")
r_n3b.font.size = Pt(8.5); r_n3b.font.color.rgb = MID_GRAY

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out_path = "/home/genmsadmin/Projects/geoasset-location/GeoAssets_Guion_Competitivo_3Slides.docx"
doc.save(out_path)
print(f"Guardado en: {out_path}")
