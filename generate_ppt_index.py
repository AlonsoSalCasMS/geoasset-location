"""
Generates GeoAssets_Intelligence_TFM_PPT_Index.docx
PPT slide-by-slide index for the GeoAssets Intelligence TFM presentation.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ── colour palette ──────────────────────────────────────────────
BLUE_DARK   = RGBColor(0x10, 0x4E, 0x8B)   # #104E8B – header/accents
BLUE_MID    = RGBColor(0x21, 0x96, 0xF3)   # #2196F3 – secondary
GREY_LIGHT  = RGBColor(0xF5, 0xF5, 0xF5)   # section bg hint
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x1A, 0x1A, 0x1A)


def set_heading_color(run, color):
    run.font.color.rgb = color


def add_slide_block(doc, slide_num: str, title: str, bullets: list[str], notes: str = ""):
    """Adds a formatted slide entry."""
    # Slide header line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)

    run_num = p.add_run(f"  Slide {slide_num}  ")
    run_num.bold = True
    run_num.font.size = Pt(10)
    run_num.font.color.rgb = WHITE
    # Shade the slide-number run
    rPr = run_num._r
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '2196F3')
    rPr.get_or_add_rPr().append(shd)

    run_title = p.add_run(f"  {title}")
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = BLUE_DARK

    # Bullet points
    for bullet in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.left_indent = Cm(1.2)
        bp.paragraph_format.space_after = Pt(1)
        run = bp.add_run(bullet)
        run.font.size = Pt(9.5)
        run.font.color.rgb = BLACK

    # Optional note
    if notes:
        np = doc.add_paragraph()
        np.paragraph_format.left_indent = Cm(1.2)
        np.paragraph_format.space_after = Pt(4)
        run_note = np.add_run(f"📌 {notes}")
        run_note.font.size = Pt(9)
        run_note.italic = True
        run_note.font.color.rgb = RGBColor(0x75, 0x75, 0x75)


def add_section_heading(doc, number: str, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

    run_num = p.add_run(f"  {number}  ")
    run_num.bold = True
    run_num.font.size = Pt(12)
    run_num.font.color.rgb = WHITE
    rPr = run_num._r
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '104E8B')
    rPr.get_or_add_rPr().append(shd)

    run_title = p.add_run(f"  {title}")
    run_title.bold = True
    run_title.font.size = Pt(13)
    run_title.font.color.rgb = BLUE_DARK

    doc.add_paragraph()  # small spacer


def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── TITLE PAGE ────────────────────────────────────────────────
    t = doc.add_heading("GeoAssets Intelligence", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.color.rgb = BLUE_DARK
        run.font.size = Pt(26)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Índice de diapositivas para la presentación del TFM")
    r.font.size = Pt(13)
    r.font.color.rgb = BLUE_MID

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub2.add_run("Máster en Inteligencia Artificial · 2025-2026")
    r2.font.size = Pt(11)
    r2.italic = True
    r2.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    doc.add_paragraph()

    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note_p.add_run(
        "Este documento describe slide por slide el contenido de la presentación.\n"
        "Cada entrada indica el título de la diapositiva, los elementos visuales/textuales "
        "recomendados y las notas de apoyo para el presentador."
    )
    nr.font.size = Pt(10)
    nr.italic = True
    nr.font.color.rgb = RGBColor(0x61, 0x61, 0x61)

    doc.add_page_break()

    # ── GLOBAL INDEX ─────────────────────────────────────────────
    idx_h = doc.add_heading("Índice de secciones", level=1)
    for run in idx_h.runs:
        run.font.color.rgb = BLUE_DARK

    sections_index = [
        ("00", "Portada y contexto de partida",        "Slides 1–2"),
        ("01", "El problema: opacidad del patrimonio empresarial", "Slides 3–7"),
        ("02", "Solución: plataforma GeoAssets Intelligence",      "Slides 8–12"),
        ("03", "Arquitectura técnica",                  "Slides 13–17"),
        ("04", "Pipelines de análisis",                 "Slides 18–27"),
        ("05", "Demostración y resultados",             "Slides 28–32"),
        ("06", "Benchmarking competitivo",              "Slides 33–35"),
        ("07", "Propuesta de valor y conclusiones",     "Slides 36–40"),
        ("08", "Bibliografía y anexos",                 "Slide 41"),
    ]

    for num, sec, slides in sections_index:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"Sección {num} – {sec}  ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BLUE_DARK
        run2 = p.add_run(slides)
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 00 – PORTADA Y CONTEXTO
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "SECCIÓN 00", "Portada y contexto de partida")

    add_slide_block(doc, "1", "Portada", [
        "Título: GeoAssets Intelligence – Localización automatizada del patrimonio productivo empresarial",
        "Subtítulo: Trabajo de Fin de Máster en Inteligencia Artificial",
        "Nombre del autor, tutor, institución y año",
        "Logo institucional + icono del proyecto (mapa con marcadores)",
    ], "Primera impresión; mantener el diseño limpio y profesional.")

    add_slide_block(doc, "2", "Punto de partida: ¿de qué trata este TFM?", [
        "Una sola frase problema: 'Localizar todos los activos productivos de una empresa requiere horas de trabajo manual sin garantía de cobertura completa'",
        "Visual: cronología de pasos manuales (buscar → contrastar → geocodificar → validar) con un reloj o símbolo de coste",
        "Transición al índice de la presentación",
    ], "Enganchar en 30 segundos; no entrar aún en detalles técnicos.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 01 – EL PROBLEMA
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "SECCIÓN 01", "El problema: opacidad del patrimonio empresarial")

    add_slide_block(doc, "3", "Contexto de mercado", [
        "Qué son los activos productivos (fábricas, oficinas, logística, retail…) y por qué importa su localización",
        "Sectores afectados: banca (riesgo colateral), seguros, consultoría estratégica, M&A, ESG",
        "Visual: mapa de España con iconos de tipos de activo",
    ])

    add_slide_block(doc, "4", "El problema actual (AS IS)", [
        "Proceso manual: búsquedas web, Registro Mercantil, LinkedIn, informes anuales",
        "Tiempo estimado: 2–4 horas por empresa para cobertura parcial",
        "Errores habituales: activos cerrados, falta de coordenadas, categorías incorrectas",
        "Visual: diagrama de flujo del proceso manual con pain points marcados en rojo",
    ], "Cifra de impacto: mencionar cuántas empresas del IBEX35 / clientes objetivo.")

    add_slide_block(doc, "5", "Limitaciones de las soluciones existentes", [
        "Fuentes públicas (CNAE, Registro): incompletas y sin geolocalización directa",
        "APIs de terceros (Google Maps): no filtran activos productivos vs. otros locales",
        "Herramientas de inteligencia competitiva: caras, no automatizadas, sin scoring",
        "Visual: tabla comparativa de limitaciones (filas = fuentes, columnas = criterios)",
    ])

    add_slide_block(doc, "6", "Oportunidad: convergencia tecnológica", [
        "LLMs capaces de razonar sobre texto + APIs geoespaciales + búsqueda web agentic",
        "Tres pilares habilitadores: Google Maps API · Modelos de lenguaje (AWS Bedrock) · Agentes IA (CrewAI)",
        "Visual: triángulo con los tres pilares, GeoAssets en el centro",
    ])

    add_slide_block(doc, "7", "Objetivo y alcance del TFM", [
        "Objetivo principal: sistema end-to-end que localiza, clasifica y puntúa activos con mínima intervención humana",
        "Alcance: empresas españolas (inicialmente); 12 categorías de activo definidas",
        "KPI de éxito: cobertura ≥ 80 % activos conocidos; tier HIGH ≥ 60 % de resultados",
        "Visual: diagrama input → sistema → output con ejemplos concretos",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 02 – LA SOLUCIÓN
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "SECCIÓN 02", "Solución: plataforma GeoAssets Intelligence")

    add_slide_block(doc, "8", "Presentación de la plataforma", [
        "Nombre, tagline: 'Inteligencia geoespacial automatizada para el patrimonio productivo'",
        "Screenshot del dashboard principal con mapa y sidebar",
        "Tres modos de análisis disponibles desde el arranque",
    ])

    add_slide_block(doc, "9", "Los tres modos de análisis", [
        "Modo 1 – Búsqueda Maps: input = nombre empresa → pipeline automático → mapa con activos",
        "Modo 2 – Carga de documento: input = PDF/Excel de activos → extracción IA → mapa",
        "Modo 3 – Agente IA: input = nombre empresa → agente web scraping → mapa",
        "Visual: tres tarjetas con icono, nombre y descripción de cada modo",
    ])

    add_slide_block(doc, "10", "Taxonomía de activos (12 categorías)", [
        "HQ, Oficina Regional, Fábrica/Planta, Centro Logístico, Centro Tecnológico",
        "Punto de Venta, Explotación Agrícola, Instalación Energética, Infraestructura Transporte",
        "Activo Hotelero, Centro Sanitario, Otro",
        "Visual: grid de iconos y colores de cada categoría (igual que en el mapa)",
    ], "Destacar que la categorización la realiza el LLM con un prompt estructurado.")

    add_slide_block(doc, "11", "Sistema de confianza: tiers y scoring", [
        "Escala 0–100 % con distribución Beta para suavizar extremos artificiales",
        "Tres tiers: HIGH (≥70 %), MEDIUM (40–69 %), LOW (<40 %)",
        "Señales por pipeline: nombre, tipo, web, reseñas B2B, valoración LLM (Maps) / evidencia documental, especificidad dirección, origen coordenadas (Documento)",
        "Visual: progress bar con los tres colores + tabla de señales y pesos",
    ])

    add_slide_block(doc, "12", "Experiencia de usuario – flujo completo", [
        "Paso 1: seleccionar modo de análisis",
        "Paso 2: introducir nombre de empresa / subir fichero / lanzar agente",
        "Paso 3: seguimiento en tiempo real con progress bar (SSE)",
        "Paso 4: mapa interactivo + sidebar con lista, filtros y exportación CSV/Excel",
        "Visual: wireframe con flechas del flujo o capturas de pantalla de cada paso",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 03 – ARQUITECTURA TÉCNICA
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "SECCIÓN 03", "Arquitectura técnica")

    add_slide_block(doc, "13", "Stack tecnológico", [
        "Frontend: Vue 3 + Vuetify 3 + Pinia + Leaflet.js",
        "Backend: FastAPI (Python 3.12) · SSE para streaming",
        "LLM: AWS Bedrock (Claude 3 Sonnet/Haiku via LiteLLM)",
        "Agente: CrewAI + DuckDuckGo MCP para búsqueda web",
        "Datos: Redis (caché) + PostgreSQL (persistencia) + Google Maps API",
        "Visual: logos del stack en dos columnas",
    ])

    add_slide_block(doc, "14", "Diagrama de arquitectura (vista global)", [
        "Tres capas: Presentación (Vue) → API Gateway (FastAPI) → Servicios (Bedrock, Maps, Redis, DB)",
        "Flujo de datos de ida (request) y vuelta (SSE stream)",
        "Indicar despliegue: contenedores Docker, opcionalmente AWS ECS",
        "Visual: diagrama de bloques con flechas y colores por capa",
    ])

    add_slide_block(doc, "15", "Modelo de datos", [
        "Entidades principales: Company, Asset",
        "Asset: 20+ campos (id, company_id, name, category, lat/lon, address, confidence_score, confidence_tier, confidence_signals, data_sources…)",
        "Persistencia: PostgreSQL (histórico) + Redis TTL 24h (caché de resultados)",
        "Visual: diagrama ER simplificado o tabla de campos con tipo y descripción",
    ])

    add_slide_block(doc, "16", "Comunicación en tiempo real con SSE", [
        "Por qué SSE (Server-Sent Events) y no WebSocket: unidireccional, nativo HTTP, sin reconexión custom",
        "Eventos del protocolo: job_started → step_start → step_complete → complete / error",
        "Frontend: EventSource + Pinia store actualizan la UI sin polling",
        "Visual: diagrama de secuencia con los eventos del SSE stream",
    ])

    add_slide_block(doc, "17", "Seguridad y configuración", [
        "Variables de entorno: GOOGLE_MAPS_API_KEY, AWS_ACCESS_KEY_ID, OPENAI_API_KEY (Bedrock proxy), DATABASE_URL, REDIS_URL",
        "Sin credenciales en código; configuración por .env + Docker secrets",
        "CORS configurado en FastAPI; API no pública (uso interno / demo)",
        "Visual: tabla de variables con tipo y descripción breve",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 04 – PIPELINES
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "SECCIÓN 04", "Pipelines de análisis")

    add_slide_block(doc, "18", "Visión general: tres pipelines", [
        "Tabla comparativa: pipeline, fuente de datos, nº pasos, tiempo estimado, mejor caso de uso",
        "Maps API → rápido, cobertura amplia, depende de datos en Google",
        "Document → muy preciso, requiere documento existente",
        "Agent IA → autónomo, sin datos previos, mayor latencia",
    ])

    # Pipeline Maps
    add_slide_block(doc, "19", "Pipeline Maps – Paso 0: Identificación de empresa", [
        "Input: nombre libre → LLM identifica empresa canónica (nombre oficial, CIF, sector, CNAE, sede)",
        "Modelo: Claude Haiku via Bedrock para bajo coste y alta velocidad",
        "Output: objeto CompanyInfo con ID normalizado",
        "Visual: ejemplo de input 'Inditex' → JSON de salida",
    ])

    add_slide_block(doc, "20", "Pipeline Maps – Pasos 1–2: Búsqueda y clasificación", [
        "Paso 1 – search_maps: múltiples queries a Google Places API con variantes del nombre; devuelve hasta N places con tipo, dirección, coordenadas",
        "Paso 2 – filter_and_classify: LLM descarta locales no productivos (franquiciados, competidores, tiendas minoristas) y asigna categoría del enum AssetCategory",
        "Visual: diagrama antes/después con ejemplos de places filtrados",
    ])

    add_slide_block(doc, "21", "Pipeline Maps – Pasos 3–4: Enriquecimiento y scoring", [
        "Paso 3 – enrich_assets: para cada activo consulta Places Details (website, phone, reviews) y extrae señales B2B",
        "Paso 4 – score_assets: calcula score ponderado de 6 señales + suavizado Beta → confidence_score, confidence_tier, confidence_signals",
        "Señales: name_match (30 %), type_match (20 %), address_corporate (15 %), website_match (15 %), reviews_b2b (10 %), llm_confidence (10 %)",
        "Visual: barra de progress con 5 pasos coloreados",
    ])

    # Pipeline Document
    add_slide_block(doc, "22", "Pipeline Documento – Pasos 0–2: Parsing y extracción", [
        "Paso 0 – parse_uploaded_document: convierte PDF/DOCX/Excel a Markdown con LlamaParse o pdfminer",
        "Paso 1 – chunk_document: divide el Markdown en chunks semánticamente coherentes",
        "Paso 2 – extract_assets_from_chunks: LLM extrae entidades de activo (nombre, dirección, tipo) de cada chunk en paralelo",
        "Visual: ejemplo de chunk de texto → JSON de activo extraído",
    ])

    add_slide_block(doc, "23", "Pipeline Documento – Pasos 3–5: Dedup, geocoding y scoring", [
        "Paso 3 – deduplicate: agrupa menciones del mismo activo por similitud de nombre/dirección",
        "Paso 4 – geocode_and_enrich: Google Geocoding API para coordenadas + Places API para enriquecimiento",
        "Paso 5 – score_document_assets: 5 señales (evidence_strength 30 %, address_specificity 20 %, coordinate_source 20 %, name_quality 15 %, llm_confidence 15 %)",
        "Visual: diagrama de los 6 pasos con nº de ítems en cada etapa (funnel)",
    ])

    # Pipeline Agent
    add_slide_block(doc, "24", "Pipeline Agente – Introducción y motivación", [
        "Caso de uso: empresa sin datos en Google Maps y sin documento disponible",
        "Enfoque: agente autónomo que navega la web como lo haría un analista humano",
        "Tecnología: CrewAI (orquestación de agentes) + DuckDuckGo MCP (búsqueda web)",
        "Visual: comparativa analista humano vs. agente IA (mismos pasos, diferente velocidad)",
    ])

    add_slide_block(doc, "25", "Pipeline Agente – Flujo de búsqueda", [
        "AgentSearchView: usuario revisa y valida ficheros descargados por el agente antes de procesar",
        "Eventos en tiempo real: thinking → searching → found_urls → downloading → accepted/rejected",
        "Criterios de aceptación de documentos: relevancia para activos productivos, formato estructurado",
        "Visual: diagrama de eventos con ejemplos de URLs aceptadas/rechazadas",
    ])

    add_slide_block(doc, "26", "Pipeline Agente – Revisión y análisis", [
        "AgentDocumentReviewView: lista de ficheros aceptados con metadata (nombre, tamaño, relevancia)",
        "El usuario confirma los documentos → se lanza el pipeline de documento sobre cada uno",
        "Procesamiento paralelo configurable (max_concurrent) para múltiples fuentes",
        "Visual: screenshot de la vista de revisión con ficheros listados",
    ])

    add_slide_block(doc, "27", "Comparativa de pipelines: cuándo usar cada uno", [
        "Árbol de decisión: ¿tienes un documento? → Documento. ¿La empresa tiene presencia en Maps? → Maps. Si no → Agente",
        "Modo combinado: posibilidad de fusionar resultados de varios pipelines (multi-source)",
        "Tabla resumen: precisión esperada, tiempo, requisitos, casos de uso ideales",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 05 – DEMOSTRACIÓN Y RESULTADOS
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "SECCIÓN 05", "Demostración y resultados")

    add_slide_block(doc, "28", "Caso de prueba 1 – Inditex (Pipeline Maps)", [
        "Input: 'Inditex'",
        "Resultados esperados: sede A Coruña, centros logísticos Arteixo/Zaragoza, tiendas insignia excluidas",
        "Métricas: nº activos encontrados, distribución por tier, tiempo de ejecución",
        "Visual: screenshot del mapa con marcadores de Inditex",
    ])

    add_slide_block(doc, "29", "Caso de prueba 2 – Repsol (Pipeline Documento)", [
        "Input: informe anual / memoria de activos de Repsol",
        "Resultados: refinerías, plantas petroquímicas, estaciones de servicio corporativas",
        "Comparativa con resultados de Maps API: diferencias de cobertura",
        "Visual: funnel con nº de activos en cada paso del pipeline",
    ])

    add_slide_block(doc, "30", "Caso de prueba 3 – Empresa sin presencia en Maps (Pipeline Agente)", [
        "Input: nombre de empresa con poca visibilidad online",
        "Logs del agente: queries ejecutadas, documentos evaluados, aceptados/rechazados",
        "Activos identificados y score de confianza obtenido",
        "Visual: captura de AgentSearchView y resultado final en mapa",
    ])

    add_slide_block(doc, "31", "Métricas de calidad del sistema", [
        "Precision@K: % de activos HIGH-tier validados manualmente como correctos",
        "Recall estimado: % de activos conocidos recuperados vs. total real",
        "Latencia por pipeline: Maps ~45s, Documento ~60s, Agente ~180s",
        "Visual: gráfico de barras o tabla comparativa de métricas por pipeline",
    ])

    add_slide_block(doc, "32", "Detalle del score de confianza en la UI", [
        "Popup de mapa: nombre, categoría, tier chip, progress bar, botón de detalle",
        "Dialog de detalle: score final, tier label, fuentes de datos (chips), señales con pesos y barras",
        "Exportación: CSV y Excel con todos los campos incluido confidence_score",
        "Visual: screenshot del dialog de detalle de confianza",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 06 – BENCHMARKING
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "SECCIÓN 06", "Benchmarking competitivo")

    add_slide_block(doc, "33", "Panorama de soluciones existentes", [
        "Categorías de competidores: herramientas de inteligencia de localización (Pitney Bowes, CARTO), plataformas de datos empresariales (Bureau van Dijk, Dun & Bradstreet), scrapers customizados",
        "Metodología de evaluación: automatización, cobertura, scoring, coste, integración",
        "Visual: mapa de posicionamiento (ejes: automatización vs. profundidad de datos)",
    ])

    add_slide_block(doc, "34", "Tabla comparativa de herramientas", [
        "Columnas: herramienta, tipo, automatización, scoring propio, precio, integración API, fuente de datos",
        "Filas: GeoAssets Intelligence, CARTO, Bureau van Dijk Orbis, Google Maps Platform (raw), solución ad-hoc con scraping",
        "Resaltar las ventajas diferenciales de GeoAssets en verde",
    ])

    add_slide_block(doc, "35", "Posicionamiento diferencial de GeoAssets", [
        "Único sistema que combina tres fuentes heterogéneas con un score unificado",
        "Open source + despliegue propio → sin dependencia de vendor, coste marginal bajo",
        "Diseñado para el caso de uso específico (activos productivos España) vs. soluciones genéricas",
        "Extensible: nuevos pipelines, nuevas categorías, nuevos modelos LLM",
        "Visual: radar chart con los 5 criterios evaluados",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 07 – PROPUESTA DE VALOR Y CONCLUSIONES
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "SECCIÓN 07", "Propuesta de valor y conclusiones")

    add_slide_block(doc, "36", "Propuesta de valor (TO BE)", [
        "Para analistas: de 4 horas a ~2 minutos por empresa con mayor cobertura",
        "Para equipos de riesgo/M&A: scoring objetivo y reproducible para due diligence",
        "Para ESG: geolocalización de activos industriales para análisis de exposición ambiental",
        "Visual: comparativa AS IS vs. TO BE en tiempo, esfuerzo y cobertura",
    ])

    add_slide_block(doc, "37", "Líneas de actuación futura", [
        "Línea 1: Expansión geográfica (Portugal, Latam) con adaptación de fuentes de datos",
        "Línea 2: Fine-tuning del modelo de clasificación con datos validados",
        "Línea 3: API pública con autenticación para integración en plataformas externas",
        "Línea 4: Módulo de detección de cambios (alertas cuando un activo desaparece o cambia)",
        "Visual: roadmap visual con fases y plazos estimados",
    ])

    add_slide_block(doc, "38", "Limitaciones actuales y mitigaciones", [
        "Dependencia de Google Maps API: datos incompletos en zonas rurales → complementar con Agente",
        "LLM no determinista: puede clasificar incorrectamente → validación por tier + revisión humana LOW",
        "Latencia del agente: ~3 minutos → uso asíncrono, notificación al completar",
        "Cobertura España: expansión futura a EMEA",
    ])

    add_slide_block(doc, "39", "Conclusiones", [
        "Validación de la hipótesis: los LLMs + APIs geoespaciales permiten automatizar el análisis de patrimonio con alta precisión",
        "Contribución técnica: sistema modular de pipelines con scoring unificado y UI interactiva",
        "Contribución académica: framework reproducible aplicable a otros dominios de inteligencia geoespacial",
        "Visual: tres bullets con icono de check en verde",
    ])

    add_slide_block(doc, "40", "Cierre y preguntas", [
        "Slide de cierre con agradecimientos",
        "QR code o URL al repositorio / demo",
        "Frase de impacto final: 'El patrimonio productivo de una empresa es visible. Solo hay que saber dónde mirar.'",
        "Contacto del autor",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 08 – BIBLIOGRAFÍA
    # ══════════════════════════════════════════════════════════════
    add_section_heading(doc, "SECCIÓN 08", "Bibliografía y anexos")

    add_slide_block(doc, "41", "Bibliografía", [
        "Google Places API Documentation – developers.google.com/maps/documentation/places",
        "LiteLLM – litellm.vercel.app (proxy LLM multi-proveedor)",
        "CrewAI Framework – crewai.com",
        "AWS Bedrock – aws.amazon.com/bedrock",
        "Brown et al. (2020) Language Models are Few-Shot Learners – GPT-3 paper",
        "Beta Distribution applied to scoring systems – referencia estadística",
        "Formato: lista numerada con estilo APA",
    ], "Completar con referencias exactas de los artículos y documentación oficial usados en el TFM.")

    # ── APPENDIX NOTES ────────────────────────────────────────────
    doc.add_page_break()
    app_h = doc.add_heading("Notas de diseño para la presentación", level=1)
    for run in app_h.runs:
        run.font.color.rgb = BLUE_DARK

    notes = [
        ("Paleta de colores", "Azul oscuro #104E8B (principal), Azul #2196F3 (secundario), Blanco, Gris claro. Colores de confianza: verde (HIGH), naranja (MEDIUM), rojo (LOW)."),
        ("Tipografía", "Montserrat para títulos (igual que la app), Inter o Calibri para cuerpo de texto. Tamaño mínimo: 18 pt en presentación."),
        ("Longitud total", "41 diapositivas (~25–30 min de presentación). Se pueden condensar las slides de detalle técnico si el tiempo es de 15 min."),
        ("Recursos visuales", "Capturas de pantalla reales de la aplicación donde se indique. Diagramas: draw.io o Mermaid exportado a PNG. Iconos: Material Design Icons (coherentes con la app)."),
        ("Notas del presentador", "Cada slide debe tener 3–5 líneas de notas en el panel inferior de PowerPoint con los puntos clave a comunicar."),
        ("Animaciones", "Mínimas: aparición de elementos por puntos en slides de lista. Sin transiciones elaboradas que distraigan."),
    ]

    for title, desc in notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run_t = p.add_run(f"{title}: ")
        run_t.bold = True
        run_t.font.size = Pt(10)
        run_t.font.color.rgb = BLUE_DARK
        run_d = p.add_run(desc)
        run_d.font.size = Pt(10)
        run_d.font.color.rgb = BLACK

    # Save
    out_path = "/home/genmsadmin/Projects/geoasset-location/GeoAssets_Intelligence_TFM_PPT_Index.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build_document()
