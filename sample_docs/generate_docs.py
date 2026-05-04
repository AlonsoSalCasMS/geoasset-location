from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_title(doc, text, color=RGBColor(0x1A, 0x56, 0x9A)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = color


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_section(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x9A)


def add_asset(doc, name, address, lat=None, lon=None, description=""):
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(address).runs[0].italic = True

    if lat and lon:
        doc.add_paragraph(f"Coordenadas: {lat}, {lon}")

    if description:
        doc.add_paragraph(description)

    doc.add_paragraph("")


# ── REPSOL ──────────────────────────────────────────────────────────────────

def build_repsol():
    doc = Document()
    doc.core_properties.author = "Repsol S.A."
    doc.core_properties.title = "Informe de Activos Productivos 2024"

    add_title(doc, "REPSOL S.A.")
    add_title(doc, "INFORME DE ACTIVOS PRODUCTIVOS 2024", RGBColor(0x22, 0x22, 0x22))
    add_subtitle(doc, "Dirección corporativa: Calle Méndez Álvaro, 44, 28045 Madrid, España")
    doc.add_paragraph("")

    add_section(doc, "1. Sede Corporativa y Oficinas Principales")

    add_asset(doc, "Sede Central Repsol",
              "Calle Méndez Álvaro, 44, 28045 Madrid, España", 40.3934, -3.6826,
              "Edificio principal del grupo con más de 3.000 empleados. Comprende las áreas de dirección general, finanzas, recursos humanos y comunicación corporativa.")

    add_asset(doc, "Campus Repsol — Ciudad de la Energía",
              "Calle Agustín de Betancourt, 4, 28003 Madrid, España", 40.4333, -3.7050,
              "Hub tecnológico y de innovación del grupo. Incluye laboratorios de I+D, centro de ciberseguridad y áreas de transformación digital.")

    add_asset(doc, "Oficina Regional Norte",
              "Calle Colón de Larreátegui, 26, 48009 Bilbao, País Vasco, España", 43.2630, -2.9350,
              "Oficina regional con funciones de coordinación operativa para el norte peninsular. Plantilla aproximada: 250 personas.")

    add_asset(doc, "Oficina Regional Sur",
              "Avenida de la Borbolla, 11, 41004 Sevilla, Andalucía, España", 37.3772, -5.9785,
              "Centro administrativo y de gestión de proyectos para la zona sur. Aproximadamente 180 empleados en funciones de operaciones y comercial.")

    add_section(doc, "2. Refinerías e Instalaciones Industriales")

    add_asset(doc, "Refinería La Pampilla — Lima, Perú (filial Repsol Perú)",
              "Avenida Néstor Gambetta 8376, Ventanilla, Callao, Perú", -11.8720, -77.1270,
              "Refinería de petróleo con capacidad de procesamiento de 117.000 barriles por día. Principal instalación productiva de Repsol en América Latina. Produce gasolinas, diésel, GLP y asfaltos.")

    add_asset(doc, "Refinería Petronor (participación Repsol 85%)",
              "Muskiz, 48550, Vizcaya, País Vasco, España", 43.3360, -3.1060,
              "Mayor refinería de España por capacidad, con 220.000 barriles/día de procesamiento. Superficie: 2,2 millones de m². Produce combustibles, lubricantes y productos petroquímicos para el mercado europeo.")

    add_asset(doc, "Refinería Puertollano",
              "Carretera de Calzada de Calatrava, s/n, 13500 Puertollano, Ciudad Real, España", 38.6860, -4.0990,
              "Capacidad de refino: 100.000 barriles/día. Especializada en producción de gasolinas de alto octanaje y combustibles de aviación. Conectada por oleoducto con el puerto de Cartagena.")

    add_asset(doc, "Refinería Cartagena",
              "Polígono Industrial El Almarjal, Escombreras, 30370 Cartagena, Murcia, España", 37.5730, -0.9810,
              "Instalación de refino ampliada en 2012. Capacidad: 220.000 barriles/día. Produce diésel de baja emisión, naftas petroquímicas y bases lubricantes. Dispone de terminal marítimo propio en el puerto de Escombreras.")

    add_asset(doc, "Complejo Petroquímico Tarragona",
              "Polígono Petroquímic, 43006 Tarragona, Cataluña, España", 41.0880, 1.2320,
              "Instalación petroquímica que produce polipropileno, polietileno y derivados. Superficie: 480 hectáreas. Emplea a más de 1.500 personas directas.")

    add_section(doc, "3. Centros Logísticos y de Distribución")

    add_asset(doc, "Terminal de Almacenamiento y Distribución — Barcelona",
              "Puerto de Barcelona, Zona Franca, Carrer B, 08040 Barcelona, España", 41.3380, 2.1550,
              "Terminal de almacenamiento de productos refinados. Capacidad total de tanques: 450.000 m³. Conectada con la red de oleoductos CLH.")

    add_asset(doc, "Terminal de Almacenamiento — Huelva",
              "Muelle de Andalucía, s/n, 21001 Huelva, Andalucía, España", 37.2650, -6.9460,
              "Instalación portuaria de descarga y almacenamiento de crudos. Capacidad: 280.000 m³.")

    add_asset(doc, "Centro Logístico Zona Centro — Madrid",
              "Calle de la Granja, 30, Polígono Industrial Pradillo, 28108 Alcobendas, Madrid, España", 40.5250, -3.6370,
              "Almacén central de distribución de lubricantes, grasas industriales y productos especiales para la zona centro peninsular. 35.000 m² de superficie cubierta.")

    add_section(doc, "4. Estaciones de Servicio (Activos Comerciales Destacados)")

    add_asset(doc, "Estación de Servicio Repsol — A-6 Km 15",
              "Autovía del Noroeste A-6, Km 15, 28220 Majadahonda, Madrid, España", 40.4590, -3.8330,
              "Estación de servicio de gran tráfico con tienda WilliamS, cafetería y lavado automático. Una de las estaciones de mayor facturación de la red en España.")

    add_asset(doc, "Estación de Servicio Repsol — Zona Franca",
              "Gran Via de les Corts Catalanes, 902, 08020 Barcelona, España", 41.4090, 2.1970,
              "Estación urbana de alta capacidad con surtidores de GNC (gas natural comprimido) y puntos de recarga eléctrica. Superficie: 4.200 m².")

    add_asset(doc, "Estación de Servicio Repsol — Sevilla Sur",
              "Carretera SE-30, Km 4,5, 41007 Sevilla, España", 37.3530, -5.9610,
              "Estación en vía de alta capacidad con área de descanso para camiones y vehículos pesados. Capacidad de almacenamiento subterráneo: 120.000 litros.")

    add_section(doc, "5. Instalaciones de Energías Renovables")

    add_asset(doc, "Parque Eólico Delta II — Aragón",
              "Término municipal de Bujaraloz, 50177 Zaragoza, Aragón, España", 41.4870, -0.1250,
              "Parque eólico con 42 aerogeneradores Vestas de 4,2 MW cada uno. Potencia total instalada: 176 MW. Producción anual estimada: 520 GWh.")

    add_asset(doc, "Planta Solar Valdesolar",
              "Término municipal de Valdecaballeros, 06683 Badajoz, Extremadura, España", 39.2540, -5.2130,
              "Planta fotovoltaica de 264 MW pico. Superficie ocupada: 750 hectáreas. En operación desde 2021, suministra energía equivalente a 185.000 hogares.")

    add_asset(doc, "Parque Eólico Offshore Wikinger (participación 25%)",
              "Mar Báltico, frente a la costa de Rügen, Alemania", 54.8330, 14.0670,
              "70 aerogeneradores offshore con potencia total de 350 MW. Repsol tiene una participación del 25% en este proyecto gestionado por Iberdrola.")

    add_section(doc, "6. Centros Tecnológicos y de Investigación")

    add_asset(doc, "Centro Tecnológico Repsol — Móstoles",
              "Calle Agustín de Betancourt, s/n, 28933 Móstoles, Madrid, España", 40.3210, -3.8730,
              "Principal centro de I+D del grupo. 500 investigadores y técnicos. Laboratorios de análisis de crudos, desarrollo de lubricantes, nuevos materiales y tecnologías de captura de CO₂. Superficie: 12.000 m².")

    add_asset(doc, "Laboratorio de Innovación Digital — Barcelona",
              "Calle Pallars, 193, 08005 Barcelona, España", 41.4010, 2.1960,
              "Hub de innovación digital enfocado en IA aplicada a operaciones, gemelos digitales de refinerías y soluciones de movilidad sostenible. 120 profesionales en plantilla.")

    doc.save("repsol_activos_productivos.docx")
    print("repsol_activos_productivos.docx generado.")


# ── INDITEX ──────────────────────────────────────────────────────────────────

def build_inditex():
    doc = Document()
    doc.core_properties.author = "Inditex S.A."
    doc.core_properties.title = "Memoria de Activos e Instalaciones Productivas 2024"

    add_title(doc, "INDITEX S.A.", RGBColor(0x00, 0x00, 0x00))
    add_title(doc, "MEMORIA DE ACTIVOS E INSTALACIONES PRODUCTIVAS 2024", RGBColor(0x22, 0x22, 0x22))
    add_subtitle(doc, "Domicilio social: Avenida de la Diputación, s/n, Edificio Inditex, 15143 Arteixo, A Coruña, España")
    doc.add_paragraph("")

    add_section(doc, "1. Sede Corporativa y Oficinas de Dirección")

    add_asset(doc, "Sede Corporativa Inditex — Arteixo",
              "Avenida de la Diputación, s/n, Edificio Inditex, 15143 Arteixo, A Coruña, Galicia, España", 43.3015, -8.5107,
              "Campus corporativo principal del grupo. Alberga la dirección general, diseño, compras, finanzas y gestión de marca. Superficie total: 200.000 m². Más de 4.500 empleados en las instalaciones centrales.")

    add_asset(doc, "Oficinas Corporativas Madrid",
              "Paseo de la Castellana, 43, 28046 Madrid, España", 40.4386, -3.6904,
              "Sede de representación institucional y relaciones con inversores. Oficinas de las áreas de comunicación corporativa, legal y relaciones públicas. Plantilla: 320 personas.")

    add_asset(doc, "Oficina Internacional Amsterdam",
              "Strawinskylaan 3105, 1077 ZX Amsterdam, Países Bajos", 52.3380, 4.8790,
              "Centro de coordinación para operaciones en Europa del Norte y mercados nórdicos. Gestiona las relaciones con proveedores europeos y distribución regional.")

    add_asset(doc, "Oficina Asia-Pacífico — Hong Kong",
              "Two Pacific Place, 88 Queensway, Admiralty, Hong Kong", 22.2760, 114.1650,
              "Hub regional para la gestión de la cadena de suministro en Asia. Coordina compras y control de calidad en Bangladesh, Vietnam, Camboya y China.")

    add_section(doc, "2. Centros Logísticos y de Distribución")

    add_asset(doc, "Plataforma Logística Central — Arteixo (Zara)",
              "Polígono Industrial de Sabón, 15142 Arteixo, A Coruña, Galicia, España", 43.2890, -8.5350,
              "Mayor centro logístico del grupo. Superficie: 500.000 m². Capacidad de procesamiento: 2,5 millones de prendas/semana. Opera 24 horas al día con sistemas automáticos de clasificación y etiquetado. Conexión directa por ferrocarril con el Puerto de A Coruña.")

    add_asset(doc, "Plataforma Logística Zaragoza",
              "Plataforma Logística de Zaragoza (PLAZA), Calle Motor, 11, 50197 Zaragoza, España", 41.6320, -0.9710,
              "Centro logístico de 130.000 m² para distribución en España, Francia, Italia y Portugal. Almacén automatizado con más de 800 estaciones robotizadas.")

    add_asset(doc, "Centro Logístico Meco — Madrid",
              "Calle de la Industria, 15, Polígono Industrial de Meco, 28880 Meco, Madrid, España", 40.5590, -3.3270,
              "Plataforma logística de 85.000 m² dedicada a Zara Home y Massimo Dutti. Gestiona distribución para toda la Península Ibérica. 1.200 empleados.")

    add_asset(doc, "Centro de Distribución Lelystad — Países Bajos",
              "Uilenkruisweg 2, 8218 PB Lelystad, Países Bajos", 52.4950, 5.4670,
              "Hub logístico europeo de 160.000 m² para distribución en Europa Central y del Norte. Capacidad de 1,8 millones de prendas procesadas por semana. Operativo desde 2023.")

    add_asset(doc, "Centro Logístico Tokyo — Japón",
              "2-1 Marunouchi, Chiyoda-ku, Tokyo 100-0005, Japón", 35.6812, 139.7671,
              "Centro de distribución para el mercado japonés y coreano. 45.000 m². Gestiona las operaciones de Zara, Pull&Bear y Bershka en Asia Oriental.")

    add_section(doc, "3. Centros de Fabricación y Talleres de Producción")

    add_asset(doc, "Planta de Confección Denllo",
              "Polígono Industrial de Ordes, Rúa das Flores, s/n, 15680 Ordes, A Coruña, Galicia, España", 43.0730, -8.4100,
              "Taller de confección propiedad del grupo (participada al 100%). Especializado en prendas de señora de alta rotación para Zara. 620 empleados, producción de 15.000 prendas/día.")

    add_asset(doc, "Centro de Producción Glencare — Galicia",
              "Polígono Industrial de Bergondo, s/n, 15165 Bergondo, A Coruña, España", 43.3470, -8.2260,
              "Instalación de patronaje, corte y confección especializada en prendas técnicas y colecciones sport. Capacidad: 10.000 unidades/día. 480 trabajadores.")

    add_asset(doc, "Taller de Estampación Fibracolor",
              "Polígono Industrial As Gándaras, Rúa Gándaras, 33, 27003 Lugo, Galicia, España", 43.0100, -7.5520,
              "Planta de estampación y tintura textil. Trata 8 toneladas de tejido/día. Certificada con ISO 14001 por su sistema de gestión ambiental.")

    add_section(doc, "4. Tiendas Insignia (Flagship Stores)")

    add_asset(doc, "Zara Flagship — Gran Vía Madrid",
              "Gran Vía, 32, 28013 Madrid, España", 40.4200, -3.7030,
              "Tienda insignia de Zara en España. Superficie: 4.500 m² distribuidos en 4 plantas. Reformada en 2022 con concepto de tienda inmersiva y probadores asistidos por IA.")

    add_asset(doc, "Zara Flagship — Passeig de Gràcia Barcelona",
              "Passeig de Gràcia, 16-18, 08007 Barcelona, España", 41.3916, 2.1694,
              "Tienda de 3.800 m² en edificio histórico modernista. Punto de venta de referencia para el mercado turístico. Dispone de servicio de click & collect.")

    add_asset(doc, "Zara Flagship — Oxford Street London",
              "118-122 Oxford Street, London W1D 1LT, Reino Unido", 51.5155, -0.1410,
              "Mayor tienda de Zara en el Reino Unido. 6.200 m² en 5 plantas. Reformada en 2023 con zona de experiencias digitales y área de moda sostenible 'Zara Join Life'.")

    add_asset(doc, "Zara Flagship — Fifth Avenue New York",
              "666 Fifth Avenue, New York, NY 10103, Estados Unidos", 40.7589, -73.9781,
              "Tienda de referencia para el mercado norteamericano. 5.400 m² en pleno Midtown Manhattan. Incluye showroom para prensa y área de personalización de prendas.")

    add_asset(doc, "Massimo Dutti — Serrano Madrid",
              "Calle Serrano, 52, 28001 Madrid, España", 40.4266, -3.6856,
              "Tienda flagship de Massimo Dutti en el eje premium de Madrid. 2.100 m² con colecciones de hombre y mujer. Servicio de sastrería a medida en planta baja.")

    add_section(doc, "5. Centros Tecnológicos y de Innovación")

    add_asset(doc, "Centro de Innovación Tecnológica — A Coruña",
              "Campus de Elviña, Universidad de A Coruña, 15071 A Coruña, España", 43.3340, -8.4110,
              "Centro de I+D+i en colaboración con la UDC. Equipos de desarrollo de IA aplicada a moda, sistemas de gestión de inventario y modelos de predicción de demanda. 80 investigadores e ingenieros.")

    add_asset(doc, "Inditex Tech Hub — Madrid",
              "Calle de Arturo Soria, 336, 28033 Madrid, España", 40.4600, -3.6330,
              "Hub tecnológico con más de 1.200 ingenieros de software trabajando en plataformas de e-commerce, integración de tienda física/digital y sistemas de RFID.")

    add_asset(doc, "Centro de Diseño Digital — Barcelona",
              "Carrer de Pallars, 98, 22@, 08018 Barcelona, España", 41.4020, 2.1980,
              "Estudio de diseño digital y creación de contenidos para todas las marcas del grupo. Equipos de fotografía, realidad aumentada, vídeo y marketing digital. 350 personas.")

    add_section(doc, "6. Edificios y Activos Inmobiliarios Corporativos")

    add_asset(doc, "Edificio Inditex Madrid — Paseo de la Castellana",
              "Paseo de la Castellana, 216, 28046 Madrid, España", 40.4660, -3.6880,
              "Edificio de oficinas de 22.000 m² alquilado para operaciones del grupo en Madrid. Certificado LEED Platinum. Alberga los equipos de e-commerce y marketing digital.")

    doc.save("inditex_activos_productivos.docx")
    print("inditex_activos_productivos.docx generado.")


build_repsol()
build_inditex()
